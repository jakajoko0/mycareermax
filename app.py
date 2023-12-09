# Import necessary libraries
import os
import stripe
from apscheduler.schedulers.background import BackgroundScheduler
from resume_builder import resume_builder  # Import the resume_builder function
from dotenv import load_dotenv
from flask import Response, stream_with_context
import pdfkit
import threading
from pathlib import Path
from io import BytesIO
import re
import time
from flask_cors import CORS, cross_origin
from flask import Flask, send_from_directory
import random
load_dotenv()
from flask_mail import Mail, Message
from itsdangerous import URLSafeTimedSerializer, SignatureExpired
import logging
#from opencensus.ext.azure.log_exporter import AzureLogHandler
#from azure.identity import DefaultAzureCredential
#from azure.keyvault.secrets import SecretClient
from docx import Document
import io
from flask import send_file
from openai import OpenAI
import requests
import docx
import pyodbc
from flask import (
    Flask,
    render_template,
    request,
    redirect,
    url_for,
    session,
    jsonify,
)
import bcrypt
from flask_login import (
    LoginManager,
    UserMixin,
    login_user,
    login_required,
    logout_user,
    current_user,
)
import urllib.parse
import json
# noqa: F401
from flask import flash, get_flashed_messages
from bs4 import BeautifulSoup
from datetime import datetime
from flask_sqlalchemy import SQLAlchemy




# Set up Azure log handler with your connection string
#logger = logging.getLogger(__name__)
#logger.addHandler(AzureLogHandler(connection_string='InstrumentationKey=24b52eea-a629-4426-816e-d819bef3b24c;IngestionEndpoint=https://eastus-8.in.applicationinsights.azure.com/;LiveEndpoint=https://eastus.livediagnostics.monitor.azure.com/'))


def create_connection(max_retries=5, wait_time_in_seconds=5):
    """Establish a connection to the Azure SQL database."""
    # Fetch database credentials from environment variables
    server = os.getenv("DB_SERVER")
    database = os.getenv("DB_NAME")
    username = os.getenv("DB_USERNAME")
    password = os.getenv("DB_PASSWORD")

    # Create the connection string
    connection_string = (
        f"DRIVER={{ODBC Driver 18 for SQL Server}};"
        f"SERVER={server};DATABASE={database};"
        f"UID={username};PWD={password}"
    )

    # Initialize variables to keep track of the connection and attempt count
    connection = None
    attempt_count = 0

    # Loop until we establish a connection or reach the maximum number of retries
    while attempt_count < max_retries:
        try:
            connection = pyodbc.connect(connection_string)
            print("Successfully connected to the database.")
            break  # Exit the loop if the connection is successful
        except pyodbc.OperationalError as e:
            print(f"Attempt {attempt_count + 1} failed with error: {e}")
            if attempt_count < max_retries - 1:
                print(f"Retrying in {wait_time_in_seconds} seconds...")
                time.sleep(wait_time_in_seconds)
            attempt_count += 1  # Increment the attempt count

    if connection is None:
        print("Max retries reached. Exiting.")
    return connection


# Example usage
if __name__ == "__main__":
    conn = create_connection()
    if conn:
        print("Connection established.")
    else:
        print("Failed to establish connection.")



# RAPIDAPI
RAPIDAPI_KEY = "04c645fbbdmshf581fe252de3b82p178cedjsn43d2da570f56"
RAPIDAPI_HOST = "jsearch.p.rapidapi.com"
RAPIDAPI_URL = "https://jsearch.p.rapidapi.com/search"

#stripe api
stripe.api_version = "2023-10-16"
stripe.api_key = "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"
#stripe webhook
endpoint_secret = 'whsec_sFYjxUvIBXWWdWUzWevSNWihQaxiDkhr'
#Logging
logger = logging.getLogger(__name__)
logging.basicConfig(level=logging.INFO)
# Create a Flask app instance and configure it
app = Flask(__name__)
app.config["DEBUG"] = True
app.config["PROPAGATE_EXCEPTIONS"] = True
app.config[
    "UPLOAD_FOLDER"
] = "/path/to/upload/folder"  # set this to your desired upload folder
app.config["SECRET_KEY"] = os.getenv("SECRET_KEY")
app.config["SESSION_COOKIE_SECURE"] = True

client = OpenAI(api_key=os.environ['OPENAI_API_KEY'])



CORS(app)


#@app.before_request
#def log_request_size():
 #   app.logger.debug(f"Request size: {len(request.get_data())}")


# Configure Flask-Mail
app.config["MAIL_SERVER"] = "smtp.office365.com"
app.config["MAIL_PORT"] = 587
app.config["MAIL_USERNAME"] = os.getenv("MAIL_USERNAME")
app.config["MAIL_PASSWORD"] = os.getenv("MAIL_PASSWORD")
app.config["MAIL_USE_TLS"] = True
app.config["MAIL_USE_SSL"] = False
mail = Mail(app)

s = URLSafeTimedSerializer(app.config["SECRET_KEY"])
node_process = None
if os.name == "nt":  # Windows
    path_wkhtmltopdf = r"C:\Program Files\wkhtmltopdf\bin\wkhtmltopdf.exe"
else:  # Assuming this is meant for non-Windows environments, such as Linux in Docker
    path_wkhtmltopdf = "/usr/bin/wkhtmltopdf"

config = pdfkit.configuration(wkhtmltopdf=path_wkhtmltopdf)
pdf_options = {"encoding": "UTF-8"}


# if os.name == "nt":  # Windows
#   path_wkhtmltopdf = "C:\\Users\\skala\\wkhtmltox-0.12.6-1.msvc2015-win64 (3).exe"
# else:  # Linux/Docker
#   path_wkhtmltopdf = "/usr/bin/wkhtmltopdf"

# config = pdfkit.configuration(wkhtmltopdf=path_wkhtmltopdf)

# pdf_options = {
#   "encoding": "UTF-8",
# }

# Configure logging

# Initialize the logger
#logger = logging.getLogger(__name__)

# Create a LoginManager instance
login_manager = LoginManager()
login_manager.login_view = "login"  # The view function that handles logins

# Initialize Flask-Login
login_manager.init_app(app)


# Create a User class (this can be an actual ORM model if you're using SQLAlchemy)
class LoginUser(UserMixin):
    def __init__(self, user_id):
        self.id = user_id


# Load the user for Flask-Login
@login_manager.user_loader
def load_user(user_id):
    return User.query.get(int(user_id))



DB_SERVER = os.getenv("DB_SERVER")
DB_NAME = os.getenv("DB_NAME")
DB_USERNAME = os.getenv("DB_USERNAME")
DB_PASSWORD = os.getenv("DB_PASSWORD")

# Update the connection string to use the fetched environment variables
conn_str = (
    "Driver={ODBC Driver 18 for SQL Server};"
    f"Server=tcp:{DB_SERVER},1433;"
    f"Database={DB_NAME};"
    f"Uid={DB_USERNAME};"
    f"Pwd={DB_PASSWORD};"
    "Encrypt=yes;"
    "TrustServerCertificate=no;"
    "Connection Timeout=90;"
)
conn = pyodbc.connect(conn_str)


def delete_user_and_associated_data(username):
    """Delete a user and their associated data from the Azure SQL database."""
    # Create database connection
    conn = create_connection()

    # Create a cursor object
    cursor = conn.cursor()

    try:
        # Get user_id for the username
        cursor.execute("SELECT id FROM dbo.users WHERE username = ?", username)
        user_id = cursor.fetchone()
        if user_id is None:
            print("Username not found.")
            return

        user_id = user_id[0]

        # SQL queries to delete associated records
        delete_documents_query = "DELETE FROM dbo.UserDocuments WHERE user_id = ?"
        delete_resumes_query = "DELETE FROM dbo.user_resumes WHERE user_id = ?"
        delete_saved_jobs_query = "DELETE FROM dbo.saved_jobs WHERE user_id = ?"

        # SQL query to delete user
        delete_user_query = "DELETE FROM dbo.users WHERE id = ?"

        # Execute the queries to delete associated records
        cursor.execute(delete_documents_query, user_id)
        cursor.execute(delete_resumes_query, user_id)
        cursor.execute(delete_saved_jobs_query, user_id)

        # Execute the query to delete the user
        cursor.execute(delete_user_query, user_id)

        # Commit the transaction
        conn.commit()
    except Exception as e:
        print(f"An error occurred: {e}")
        conn.rollback()
    finally:
        # Close the connection
        conn.close()


@app.route("/extension_test")
def extension_test():
    return render_template("extension_test.html")


@app.route("/submitted")
def submitted():
    return render_template("submitted.html")

@app.route("/purchase")
#@login_required 
def purchase():
    user_id = current_user.id if current_user.is_authenticated else None
    return render_template("purchase.html", user_id=user_id)

@app.route("/subscription")
def subscription():
    user_id = current_user.id if current_user.is_authenticated else None
    return render_template("subscription.html", user_id=user_id)    

@app.route("/checkout")
@login_required 
def checkout():
    user_id = current_user.id if current_user.is_authenticated else None
    return render_template("checkout.html", user_id=user_id)

@app.route('/success', methods=['GET'])
def success():
    user_id = current_user.id if current_user.is_authenticated else None
    return render_template("success.html", user_id=user_id)

@app.route('/cancel', methods=['GET'])
def cancel():
    user_id = current_user.id if current_user.is_authenticated else None
    return render_template("cancel.html", user_id=user_id)


def retrieve_latest_resume(user_id):
    try:
        with create_connection().cursor() as cursor:
            query = "SELECT TOP 1 resume_text FROM user_resumes WHERE user_id = ? ORDER BY uploaded_at DESC"
            cursor.execute(query, (user_id,))
            row = cursor.fetchone()
            if row:
                return row[0]
            else:
                return None
    except Exception as e:
        # In a production environment, consider logging the exception for debugging.
        return None

@app.route("/get_resume_max", methods=["GET"])
def get_resume_max():
    if current_user.is_authenticated:
        user_id = current_user.get_id()
        resume_content = retrieve_latest_resume(user_id)
        if resume_content:
            return jsonify({"success": True, "resume_content": resume_content})
        else:
            return jsonify({"success": False})
    else:
        return jsonify({"success": False})


@app.route("/googleba8e248f290d00cf.html")
def google_verification():
    return render_template("googleba8e248f290d00cf.html")


@app.route("/app-ads.txt")
def serve_app_ads_txt():
    return send_from_directory("static", "app-ads.txt")


@app.route("/delete_account", methods=["GET", "POST"])
def delete_account():
    if request.method == "POST":
        username = request.form["username"]
        delete_user_and_associated_data(username)
        flash("Account successfully deleted.")
        return redirect(url_for("delete_account"))
    return render_template("delete_account.html")


@app.route("/forgot-password", methods=["GET", "POST"])
def forgot_password():
    if request.method == "POST":
        email = request.form["email"]

        # Generate token
        token = s.dumps(email, salt="email-confirm")

        # Create email message
        msg = Message(
            "myCAREERMAX Password Reset Request",
            sender="password_reset@mycareermax.com",
            recipients=[email],
        )
        link = url_for("reset_token", token=token, _external=True)

        msg.body = (
            f"Please use this link to reset your password: {link}\n\n"
            "---\n"
            "Note: This inbox is not monitored. Please do not reply to this email. "
            "If you have any questions or need further assistance, please contact us at "
            "support@mycareermax.com"
        )

        # Send email
        mail.send(msg)

        return "Email has been sent!"

    return render_template("forgot_password.html")


@app.route("/reset-password/<token>", methods=["GET", "POST"])
def reset_token(token):
    try:
        email = s.loads(token, salt="email-confirm", max_age=3600)
    except SignatureExpired:
        return "The token is expired!"

    if request.method == "POST":
        new_password = request.form["new_password"]

        # Hash the new password
        hashed_password = bcrypt.hashpw(new_password.encode("utf-8"), bcrypt.gensalt())

        # Update the password in the database
        with pyodbc.connect(conn_str) as conn:
            cursor = conn.cursor()
            cursor.execute(
                "UPDATE users SET password = ? WHERE email = ?",
                (hashed_password, email),
            )
            conn.commit()

        login_url = url_for("login")
        return f'Password Reset Successfully! <a href="{login_url}">Login</a>'

    return render_template("reset_token.html")

@app.route("/register", methods=["GET", "POST"])
def register():
    message = None
    if request.method == "POST":
        username = request.form["username"]
        email = request.form["email"]
        password = request.form["password"]
        confirm_password = request.form["confirm_password"]

        # Basic validation
        if not username or not email or not password or not confirm_password:
            message = "All fields are required!"
        elif password != confirm_password:
            message = "Passwords do not match!"
        else:
            with pyodbc.connect(conn_str) as conn:
                cursor = conn.cursor()
                cursor.execute("SELECT * FROM users WHERE username = ?", (username,))
                user = cursor.fetchone()
                if user:
                    message = "Username already exists!"
                else:
                    # Hash the password
                    hashed_password = bcrypt.hashpw(password.encode("utf-8"), bcrypt.gensalt())

                    # Insert user data into the database (without Stripe customer ID)
                    cursor.execute(
                        "INSERT INTO users (username, email, password) VALUES (?, ?, ?)",
                        (username, email, hashed_password)
                    )
                    conn.commit()
                    message = "User registered successfully!"
                    return redirect(url_for("login"))

    return render_template("register.html", message=message)



@app.route("/login", methods=["GET", "POST"])
def login():
    message = None
    if request.method == "POST":
        username = request.form["username"]
        password_to_check = request.form["password"].encode("utf-8")
        remember_me = bool(request.form.get("remember_me"))

        cursor = conn.cursor()
        cursor.execute("SELECT * FROM users WHERE username = ?", (username,))
        user_record = cursor.fetchone()

        # Decode the hashed password from the database before comparing
        if user_record and bcrypt.checkpw(password_to_check, user_record.password):
            # Create a user object without using SQLAlchemy
            user_obj = LoginUser(user_record.id)
            login_user(user_obj, remember=remember_me)
            return redirect(url_for("dashboard"))
        else:
            message = "Invalid credentials"

    return render_template("login.html", message=message)


db = SQLAlchemy()

# Quote the username and password to handle special characters
quoted_username = urllib.parse.quote_plus(os.environ["DB_USERNAME"])
quoted_password = urllib.parse.quote_plus(os.environ["DB_PASSWORD"])

# Construct the connection string
connection_string = (
    f"mssql+pyodbc://{quoted_username}:{quoted_password}@{os.environ['DB_SERVER']}:1433/"
    f"{os.environ['DB_NAME']}?driver={urllib.parse.quote_plus('ODBC Driver 18 for SQL Server')}"
)

app.config["SQLALCHEMY_DATABASE_URI"] = connection_string
db.init_app(app)

from sqlalchemy import ForeignKey
from sqlalchemy import desc

# Filter for "Month Year" format
@app.template_filter('format_date_month_year')
def format_date_month_year(d):
    return d.strftime('%B %Y') if d else "Present"  # e.g., "May 2020" or "Present"

# Filter for "MM/YYYY" format
@app.template_filter('format_date_numeric')
def format_date_numeric(d):
    return d.strftime('%m/%Y') if d else "Present"  # e.g., "05/2020" or "Present"

class UserResumes(db.Model):
    __tablename__ = 'user_resumes'
    id = db.Column(db.Integer, primary_key=True)
    user_id = db.Column(db.Integer, db.ForeignKey('users.id'), nullable=False)
    resume_text = db.Column(db.Text)
    filename = db.Column(db.String(255))
    uploaded_at = db.Column(db.DateTime, default=datetime.utcnow)
    questionnaire = db.Column(db.Text)

    # Define the relationship
    user = db.relationship('User', backref=db.backref('resumes', lazy=True))

    def __init__(self, user_id, resume_text, filename=None, questionnaire=None):
        self.user_id = user_id
        self.resume_text = resume_text
        self.filename = filename
        self.uploaded_at = datetime.utcnow()
        self.questionnaire = questionnaire



class User(db.Model):
    __tablename__ = 'users'
    id = db.Column(db.Integer, primary_key=True)
    username = db.Column(db.String(255), nullable=False, unique=True)
    password = db.Column(db.LargeBinary, nullable=False)
    email = db.Column(db.String(255), unique=True)
    created_at = db.Column(db.DateTime, default=datetime.utcnow)
    access_token = db.Column(db.String(255))
    stripe_customer_id = db.Column(db.String(255))


    @property
    def is_authenticated(self):
        # Always return True as all users are authenticated by default
        return True

    @property
    def is_active(self):
        # Return True if this is an active user
        return True

    @property
    def is_anonymous(self):
        # Return False as anonymous users aren't supported
        return False

    def get_id(self):
        # Return the user id as a unicode string
        return str(self.id)


class UserDocuments(db.Model):
    __tablename__ = "UserDocuments"

    id = db.Column(db.Integer, primary_key=True, autoincrement=True)
    user_id = db.Column(db.Integer, db.ForeignKey("users.id"), nullable=False)
    document_name = db.Column(db.String(255), nullable=False)
    document_content = db.Column(
        db.String, nullable=False
    )  # Consider changing to db.Text or db.LargeBinary if the content can be large
    document_type = db.Column(db.String(50), nullable=False)
    creation_date = db.Column(db.DateTime, default=datetime.utcnow)
    job_title = db.Column(db.String(255), nullable=True)  # Add job title
    company_name = db.Column(db.String(255), nullable=True)  # Add company name
    apply_link = db.Column(db.String(255), nullable=True)  # Add apply link
    job_id = db.Column(db.Integer, nullable=True)  # or nullable=False based on your requirements

    user = db.relationship("User", backref=db.backref("documents", lazy=True))




class Summary(db.Model):
    __tablename__ = 'summary'
    id = db.Column(db.Integer, primary_key=True)
    user_id = db.Column(db.Integer, db.ForeignKey('users.id'))
    summary_text = db.Column(db.Text)
    created_at = db.Column(db.DateTime, default=datetime.utcnow, nullable=False)  # Add the created_at field



class Education(db.Model):
    __tablename__ = 'education'
    id = db.Column(db.Integer, primary_key=True)
    user_id = db.Column(db.Integer, db.ForeignKey('users.id'))
    degree = db.Column(db.String(255))
    institution = db.Column(db.String(255))
    graduation_year = db.Column(db.Integer)



class Skills(db.Model):
    __tablename__ = 'skills'
    id = db.Column(db.Integer, primary_key=True)
    user_id = db.Column(db.Integer, db.ForeignKey('users.id'))
    skill = db.Column(db.String(255))

class Projects(db.Model):
    __tablename__ = 'projects'
    id = db.Column(db.Integer, primary_key=True)
    user_id = db.Column(db.Integer, db.ForeignKey('users.id'))
    project_name = db.Column(db.String(255))
    description = db.Column(db.Text)
    url = db.Column(db.String(255))

class ProjectBulletPoint(db.Model):
    __tablename__ = 'projects_bullet_points'
    id = db.Column(db.Integer, primary_key=True)
    project_id = db.Column(db.Integer, db.ForeignKey('projects.id'), nullable=False) 
    text = db.Column(db.String(255))

    project = db.relationship('Projects', backref=db.backref('bullet_points', lazy=True)) 


class PersonalInformation(db.Model):
    __tablename__ = 'personal_information'
    id = db.Column(db.Integer, primary_key=True)
    user_id = db.Column(db.Integer, db.ForeignKey('users.id'))
    full_name = db.Column(db.String(255))
    email = db.Column(db.String(255))
    contact_number = db.Column(db.String(50))
    city_of_residence = db.Column(db.String(255))
    state_of_residence = db.Column(db.String(255))
    website = db.Column(db.String(255))
    github = db.Column(db.String(255))
    linkedin = db.Column(db.String(255))

    # Define the relationship
    user = db.relationship('User', backref=db.backref('personal_info', uselist=False, lazy=True))


class Certifications(db.Model):
    __tablename__ = 'certifications'
    id = db.Column(db.Integer, primary_key=True)
    user_id = db.Column(db.Integer, db.ForeignKey('users.id'))
    certification_name = db.Column(db.String(255))
    issued_by = db.Column(db.String(255))
    issue_date = db.Column(db.Date)
    expiration_date = db.Column(db.Date)
    certification_url = db.Column(db.String(255))

class WorkExperienceBulletPoint(db.Model):
    __tablename__ = 'work_experience_bullet_points'
    id = db.Column(db.Integer, primary_key=True)
    work_experience_id = db.Column(db.Integer, db.ForeignKey('work_experience.id'), nullable=False)
    text = db.Column(db.String(255))

    work_experience = db.relationship('WorkExperience', backref=db.backref('bullet_points', lazy=True))

class WorkExperience(db.Model):
    __tablename__ = 'work_experience'
    id = db.Column(db.Integer, primary_key=True)
    user_id = db.Column(db.Integer, db.ForeignKey('users.id'))
    job_title = db.Column(db.String(255))
    company = db.Column(db.String(255))
    start_date = db.Column(db.Date)
    end_date = db.Column(db.Date)

class UserDetails(db.Model):
    __tablename__ = 'UserDetails'
    UserID = db.Column(db.Integer, db.ForeignKey('users.id'), primary_key=True)
    FullName = db.Column(db.String(255))
    RecentPosition = db.Column(db.String(255))
    DesiredJobTitle = db.Column(db.String(255))
    DesiredJobLocation = db.Column(db.String(255))
    DesiredWorkType = db.Column(db.String(50))
    DesiredCompensation = db.Column(db.String(50))
    JobAlertNotifications = db.Column(db.String(3))

class Subscription(db.Model):
    __tablename__ = 'subscriptions'
    id = db.Column(db.Integer, primary_key=True)
    user_id = db.Column(db.Integer, db.ForeignKey('users.id'), unique=True)
    tier = db.Column(db.String(50))
    start_date = db.Column(db.DateTime, default=datetime.utcnow)
    end_date = db.Column(db.DateTime)

    user = db.relationship('User', backref=db.backref('subscription', uselist=False))

    def is_active(self):
        return self.end_date is None or datetime.utcnow() <= self.end_date



# Reverse mapping from Price IDs to Tier Names
tier_from_price_id = {
    'price_1OJOPqBmOXAq5RyDWYawmWvf': 'Tier1',
    'price_1OJOPyBmOXAq5RyDzIAxE0tc': 'Tier2',
    'price_1OJXflBmOXAq5RyDxEqGCZMg': 'FreePlan'
}

@app.route('/webhook', methods=['POST'])
def webhook():
    payload = request.get_data(as_text=True)
    sig_header = request.headers['STRIPE_SIGNATURE']
    event = None
    try:
        event = stripe.Webhook.construct_event(payload, sig_header, endpoint_secret)
        data = event['data']
        event_type = event['type']
        data_object = data['object']

        logger.info(f'Received Stripe webhook event: {event_type}')

        if event_type in ['customer.subscription.created', 'customer.subscription.updated', 'customer.subscription.deleted']:
            stripe_customer_id = data_object['customer']

            # Retrieve the customer's email from Stripe
            stripe_customer = stripe.Customer.retrieve(stripe_customer_id)
            customer_email = stripe_customer['email']

            conn = create_connection()
            cursor = conn.cursor()

            # Update or Insert user with the new stripe_customer_id
            cursor.execute("""
                IF EXISTS (SELECT 1 FROM dbo.users WHERE email = ?)
                    UPDATE dbo.users SET stripe_customer_id = ? WHERE email = ?
                ELSE
                    INSERT INTO dbo.users (email, stripe_customer_id) VALUES (?, ?)
                """, customer_email, stripe_customer_id, customer_email, customer_email, stripe_customer_id)

            # Fetch User ID using email address
            cursor.execute("SELECT id FROM dbo.users WHERE email = ?", customer_email)
            user_id = cursor.fetchone()[0]

            logger.info(f'User ID {user_id} updated with Stripe customer ID: {stripe_customer_id}')

            if event_type in ['customer.subscription.created', 'customer.subscription.updated']:
                price_id = data_object['items']['data'][0]['price']['id']
                tier = tier_from_price_id.get(price_id, 'Unknown')
                start_date = datetime.fromtimestamp(data_object['current_period_start'])
                end_date = datetime.fromtimestamp(data_object['current_period_end'])

                # Check for existing subscription
                cursor.execute("SELECT id FROM dbo.subscriptions WHERE user_id = ?", user_id)
                subscription_result = cursor.fetchone()

                if subscription_result:
                    # Update existing subscription
                    logger.info(f'Updating subscription: User ID {user_id}, Tier {tier}')
                    cursor.execute(
                        "UPDATE dbo.subscriptions SET tier = ?, start_date = ?, end_date = ? WHERE user_id = ?",
                        tier, start_date, end_date, user_id
                    )
                else:
                    # Create new subscription
                    logger.info(f'Creating new subscription: User ID {user_id}, Tier {tier}')
                    cursor.execute(
                        "INSERT INTO dbo.subscriptions (user_id, tier, start_date, end_date) VALUES (?, ?, ?, ?)",
                        user_id, tier, start_date, end_date
                    )

            elif event_type == 'customer.subscription.deleted':
                logger.info(f'Processing subscription deletion for User ID {user_id}')
                # Set the end date for all subscriptions of the user to now
                current_time = datetime.utcnow()
                cursor.execute(
                    "UPDATE dbo.subscriptions SET end_date = ? WHERE user_id = ? AND (end_date > ? OR end_date IS NULL)",
                    current_time, user_id, current_time
                )

            conn.commit()
            conn.close()

    except stripe.error.SignatureVerificationError as e:
        logger.error(f'Signature verification failed: {e}')
        return jsonify({'error': 'Invalid signature'}), 400
    except Exception as e:
        logger.exception(f'Internal server error: {e}')
        return jsonify({'error': 'Internal server error'}), 500

    logger.info('Stripe webhook processed successfully')
    return jsonify({'status': 'success'})



@app.route('/.well-known/<path:filename>')
def well_known(filename):
    directory = os.path.join(app.root_path, '.well-known')
    return send_from_directory(directory, filename)


@app.route("/get-username")
@login_required
def get_username():
    try:
        with conn.cursor() as cursor:
            query = "SELECT username FROM users WHERE id = ?"
            cursor.execute(query, (current_user.id,))
            result = cursor.fetchone()
            if result:
                username = result[0]
                return jsonify({"success": True, "username": username})
            else:
                return jsonify({"success": False, "error": "User not found."})
    except Exception as e:
       # app.logger.error(f"Failed to fetch username: {e}")
        return jsonify({"success": False, "error": "Failed to fetch username."})


@app.route("/save_document", methods=["POST"])
def save_document():
    try:
        # Get the request data from JSON (since we're sending JSON from the client-side)
        data = request.json
        user_id = data.get("user_id")
        doc_name = data.get("document_name")
        doc_content = data.get("document_content")
        doc_type = data.get("document_type")
        job_title = data.get("job_title")
        company_name = data.get("company_name")
        apply_link = data.get("apply_link")

        # Create the new document object with all the details
        new_document = UserDocuments(
            user_id=user_id,
            document_name=doc_name,
            document_content=doc_content,
            document_type=doc_type,
            job_title=job_title,  # Include job title
            company_name=company_name,  # Include company name
            apply_link=apply_link,  # Include apply link
        )

        # Add and commit the new document
        db.session.add(new_document)
        db.session.commit()
        return jsonify(
            {
                "message": "Document saved successfully. You may view your Saved Documents on your Dashboard",
                "doc_id": new_document.id,
            }
        )
    except Exception as e:
        db.session.rollback()
        return (
            jsonify(
                {
                    "error": "An error occurred while saving the document: {}".format(
                        str(e)
                    )
                }
            ),
            500,
        )


# @app.route("/dashboard", methods=["GET"])
# @login_required
# def get_user_documents():
#   user_id = current_user.id
#  user_documents = UserDocuments.query.filter_by(user_id=user_id).all()

# documents_list = [
#   {
#        "id": doc.id,
#      "document_name": doc.document_name,
#     "document_content": doc.document_content,
#    "document_type": doc.document_type,
#   "creation_date": doc.creation_date,
#  "job_title": doc.job_title,  # Include job title
# "company_name": doc.company_name,  # Include company name
# "apply_link": doc.apply_link,  # Include apply link
#    }
#   for doc in user_documents
#  ]

# return render_template("dashboard.html", documents=documents_list)

@app.route("/delete_document/<int:document_id>", methods=["DELETE"])
@login_required
def delete_document(document_id):
    try:
        doc = UserDocuments.query.filter_by(id=document_id).first()
        if not doc:
            return jsonify({"success": False, "error": "Document not found"}), 404

        db.session.delete(doc)
        db.session.commit()
        return jsonify({"success": True, "message": "Document deleted successfully"})
    except Exception as e:
        db.session.rollback()
        return jsonify({"success": False, "error": str(e)}), 500

def analyze_resume_compatibility(job_description, user_resume):
    response = client.chat.completions.create(
        model="gpt-3.5-turbo-1106",
        messages=[
            {
                "role": "system",
                "content": "You are an expert at analyzing resumes and job descriptions."
            },
            {
                "role": "user",
                "content": f"Given the job description: '{job_description}' and user resume: '{user_resume}', provide a summarized job description, a job compatibility percentage, and an extracted list of keywords from the resume and job description in the following format:\n\n1. keywords_job: list keywords extracted separated by commas.\n\n2. Compatibility %: x%\n\n3. Job Summary: write your brief summary of the job description here.\n\n4. keyword_resume: list keyword extracted from resume here."
            }
        ],
        temperature=0.6,
        # max_tokens=3711,
        top_p=1,
        frequency_penalty=0,
        presence_penalty=0
    )

    chatgpt_response = response.choices[0].message.content
    logging.debug(f"ChatGPT Response: {chatgpt_response}")

    return chatgpt_response


# Example usage:
#job_description_input = "Your provided job description here..."
#user_resume_input = "Your provided resume here..."
#result = analyze_resume_compatibility(job_description_input, user_resume_input)
#print(result)

def extract_compatibility(chatgpt_response):
    compatibility_match = re.search(r'(compatibility|percentage).*?(\d+)%', chatgpt_response, re.IGNORECASE)
    if compatibility_match:
        return float(compatibility_match.group(2))
    return None

def extract_summary(chatgpt_response):
    description_match = re.search(r'Job Summary: (.*?)(4\. keyword_resume|$)', chatgpt_response, re.IGNORECASE | re.DOTALL)
    if description_match:
        return description_match.group(1).strip()
    return None

def extract_keywords_resume(chatgpt_response):
    keywords_match = re.search(r'(\b\w*resume\w*\b).*?:\s*(.+?)(?=\b\w*job\w*\b|\b\w*sum\w*\b|compatibility|$)', chatgpt_response, re.IGNORECASE | re.DOTALL)
    if keywords_match:
        keywords_str = keywords_match.group(2)
        return [keyword.strip() for keyword in re.split(',|\n', keywords_str)]
    return []

def extract_keywords_job(chatgpt_response):
    keywords_match = re.search(r'(\b\w*job\w*\b).*?:\s*(.+?)(?=\b\w*resume\w*\b|\b\w*sum\w*\b|compatibility|$)', chatgpt_response, re.IGNORECASE | re.DOTALL)
    if keywords_match:
        keywords_str = keywords_match.group(2)
        return [keyword.strip() for keyword in re.split(',|\n', keywords_str)]
    return []




@app.route("/save-job", methods=["POST"])
@login_required
def save_job():
    logging.info("Received request to /save-job.")
    try:
        data = request.get_json()
        logging.debug(f"Request data: {data}")
        
        user_id = current_user.id
        logging.debug(f"Current user ID: {user_id}")
        
        job_id = data.get("job_id")
        job_title = data.get("job_title")
        original_job_description = data.get("job_description")
        job_link = data.get("job_link")
        job_loc = data.get("job_loc")
        company_name = data.get("company_name")
        link = data.get("link")
        employer_logo = data.get("employer_logo")

        user_resume = retrieve_latest_resume(user_id)
        logging.debug(f"Retrieved user resume: {user_resume}")

        chatgpt_response = analyze_resume_compatibility(original_job_description, user_resume)
        logging.debug(f"Received ChatGPT response: {chatgpt_response}")

        compatibility_score = extract_compatibility(chatgpt_response)
        keywords_resume = extract_keywords_resume(chatgpt_response)
        keywords_job = extract_keywords_job(chatgpt_response)
        job_description = extract_summary(chatgpt_response)  

        logging.debug(f"Extracted values: compatibility_score={compatibility_score}, keywords_resume={keywords_resume}, keywords_job={keywords_job}, job_description={job_description}")

        with conn.cursor() as cursor:
            cursor.execute(
                """
                INSERT INTO saved_jobs (user_id, job_id, job_title, job_description, job_link, job_loc, 
                                        company_name, link, employer_logo, percentage, keywords_resume, keywords_job) 
                VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
                """,
                (
                    user_id,
                    job_id,
                    job_title,
                    job_description,
                    job_link,
                    job_loc,
                    company_name,
                    link,
                    employer_logo,
                    compatibility_score,
                    ",".join(keywords_resume),
                    ",".join(keywords_job),
                ),
            )
            conn.commit()
            logging.info("Job saved successfully to the database.")

        return jsonify({"success": True, "message": "Job saved successfully!"})

    except Exception as e:
        logging.error(f"Error encountered: {str(e)}")
        return jsonify({"success": False, "error": str(e)}), 500



# GET SAVED JOBS
@app.route("/get-saved-jobs", methods=["GET"])
@login_required
def get_saved_jobs():
    logging.info("Received request to /get-saved-jobs.")
    try:
        user_id = current_user.id
        logging.debug(f"Current user ID: {user_id}")
        
        # Create a new connection for this route
        local_conn = create_connection()
        logging.debug("Database connection established.")
        
        with local_conn.cursor() as cursor:
            cursor.execute("SELECT * FROM saved_jobs WHERE user_id = ?", (user_id,))
            saved_jobs = cursor.fetchall()
        logging.debug(f"Retrieved {len(saved_jobs)} saved jobs from the database.")

        # Convert the saved_jobs data to a list of dictionaries
        saved_jobs_list = []
        for job in saved_jobs:
            job_dict = {
                "id": job.id,
                "job_id": job.job_id,
                "job_title": job.job_title,
                "job_description": job.job_description,
                "job_link": job.job_link,
                "job_loc": job.job_loc,
                "company_name": job.company_name,
                "link": job.link,
                "employer_logo": job.employer_logo,
                "status": job.status,  # Add the status field here
                "compatibility": job.percentage  # Add the compatibility field here
            }
            saved_jobs_list.append(job_dict)
        logging.debug(f"Converted saved jobs to a list of dictionaries.")

        local_conn.close()  # Close the local connection
        logging.info("Database connection closed.")
        
        return jsonify({"success": True, "saved_jobs": saved_jobs_list})

    except Exception as e:
        logging.error(f"Error encountered: {str(e)}")
        return jsonify({"success": False, "error": str(e)}), 500

@app.route("/remove-saved-job", methods=["POST"])
@login_required
def remove_saved_job():
    try:
        data = request.get_json()
        job_to_remove_id = data.get("job_id")
        user_id = current_user.id

        with conn.cursor() as cursor:
            cursor.execute(
                "DELETE FROM saved_jobs WHERE user_id = ? AND id = ?",
                (user_id, job_to_remove_id),
            )
            conn.commit()

        return jsonify({"success": True, "message": "Job removed successfully!"})
    except Exception as e:
        return jsonify({"success": False, "error": str(e)}), 500


@app.route("/upload_dashresume", methods=["POST"])
@login_required
def upload_dashresume():
    if "resume" not in request.files:
        return jsonify({"success": False, "error": "No file part"})
    file = request.files["resume"]

    if file.filename == "":
        return jsonify({"success": False, "error": "No selected file"})

    if file and file.filename.endswith(".docx"):
        document = Document(io.BytesIO(file.read()))
        resume_text = "\n".join([p.text for p in document.paragraphs if p.text])

        try:
            with conn.cursor() as cursor:
                # Assuming you want to overwrite the old resume:
                cursor.execute(
                    "DELETE FROM user_resumes WHERE user_id = ?", (current_user.id,)
                )
                cursor.execute(
                    "INSERT INTO user_resumes (user_id, resume_text, filename) VALUES (?, ?, ?)",
                    (current_user.id, resume_text, file.filename),
                )
                conn.commit()
        except Exception as e:
        #   app.logger.error(f"Failed to insert resume into database: {e}")
            return jsonify({"success": False, "error": "Failed to upload resume."})
        return jsonify({"success": True, "filename": file.filename})
    return jsonify({"success": False, "error": "Invalid file format"})


@app.route("/get_latest_resume_name")
@login_required
def get_latest_resume_name():
    print("Fetching latest resume name...")
    try:
        with create_connection().cursor() as cursor:  # Create a new connection
            query = "SELECT TOP (1) filename FROM user_resumes WHERE user_id = ? ORDER BY uploaded_at DESC"
            cursor.execute(query, (current_user.id,))
            result = cursor.fetchone()
            if result:
                return jsonify({"success": True, "filename": result[0]})
            else:
                return jsonify({"success": False, "error": "No resume found."})
    except Exception as e:
    #    app.logger.error(f"Failed to fetch resume name: {e}")
        return jsonify({"success": False, "error": "Failed to fetch resume name."})

        # DELETE RESUME FROM DASHBOARD


@app.route("/delete_resume", methods=["POST"])
@login_required
def delete_resume():
    try:
        with conn.cursor() as cursor:
            cursor.execute(
                "DELETE FROM user_resumes WHERE user_id = ?", (current_user.id,)
            )
            conn.commit()
            return jsonify({"success": True})
    except Exception as e:
    #    app.logger.error(f"Failed to delete resume: {e}")
        return jsonify({"success": False, "error": "Failed to delete resume."})


@app.route("/download_document/<int:document_id>", methods=["GET"])
def download_document(document_id):
    # Fetch the document from the database
    doc = UserDocuments.query.filter_by(id=document_id).first()
    if not doc:
        return "Document not found", 404

    # Determine the format for download (default is docx)
    download_format = request.args.get("format", "docx")

    if download_format == "docx":
        # Create a new Document
        document = Document()
        document.add_paragraph(doc.document_content)

        # Stream the .docx file back to the user
        output = io.BytesIO()
        document.save(output)
        output.seek(0)
        return send_file(
            output,
            as_attachment=True,
            download_name=f"{doc.document_name}.docx",
            mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )

    elif download_format == "pdf":
        # Convert the document content to basic HTML with UTF-8 meta tag
        html_content = """<!DOCTYPE html>
        <html>
        <head>
            <meta charset="UTF-8">
        </head>
        <body>
        <p>{}</p>
        </body>
        </html>""".format(
            doc.document_content.replace("\n", "<br>")
        )

        # Convert the HTML to PDF using pdfkit
        pdf_content = pdfkit.from_string(
            html_content, False, configuration=config, options=pdf_options
        )

        # Stream the PDF back to the user without a specific filename
        response = Response(stream_with_context(io.BytesIO(pdf_content)))
        response.headers["Content-Type"] = "application/pdf"
        response.headers["Content-Disposition"] = "inline"
        return response

    else:
        return "Invalid format specified", 400

@app.route("/download_html2pdf/<int:document_id>", methods=["GET"])
def download_html2pdf(document_id):
    # Fetch the document from the database
    doc = UserDocuments.query.filter_by(id=document_id).first()
    if not doc:
        return "Document not found", 404

    download_format = request.args.get("format", "docx")

    if download_format == "pdf":
        # Define the HTML head with styles
        html_head = """
        <html>
        <head>
            <title>Resume</title>
            <meta charset="UTF-8">
            <style>
   /*------------------------GLOBAL STYLES --------------------------*/
        body {
            font-size: 12px;
            /* Base font size */
        }

        p,
        li,
        a {
            font-size: 1em;
            /* Relative to base font size */
        }

        /* Section Title Headers */
        h2 {
            font-size: 2em;
            text-align: center;
            margin-bottom: 5px;
        }

        hr {
            border: none;
            height: 1px;
            background-color: #c4c4c4;
            margin-top: 10px;
            margin-bottom: 10px;
            margin-left: 10px;
            margin-right: 10px;
            width: 100%;
        }

        @media print {
            hr {
                /* Specific style for print media */
                background-color: #c4c4c4;
            }
        }

        /*------------------------PERSONAL INFORMATION --------------------------*/
        .personal-information .personal-name {
            font-size: 3em;
        }

        .personal-information {
            text-align: center;
        }

        .personal-details {
            margin-top: 0;
            margin-bottom: 2px;
        }

        /*------------------------SUMMARY SECTION --------------------------*/
        .summary-section p {
            margin-top: 5px;
            margin-bottom: 5px;
        }

        /*------------------------WORK EXPERIENCE --------------------------*/
        .work-experience-entry .title-dates {
            display: flex;
            justify-content: space-between;
            align-items: center;
        }

        .job-title {
            font-weight: bold;
            font-size: 1.1em;
        }

        .dates {
            font-size: 1em;
            margin-left: auto;
            font-weight: bold;
        }

        .company-name {
            font-size: 1em;
        }

        .bullet-points {
            font-size: 0.9em;
        }


        /*------------------------EDUCATION --------------------------*/
        .education-section p {
            margin-top: 5px;
            margin-bottom: 5px;
        }

        /*------------------------ILLS --------------------------*/
        .skills-list {
            list-style-type: none;
            padding: 0;
            font-size: 1em;
        }

        .skills-list li {
            display: inline;
            margin-right: 10px;
        }

        .skills-list li::after {
            content: '•';
            margin-left: 10px;
        }

        .skills-list li:last-child::after {
            content: '';
        }

        /*------------------------CERTIFICATIONS SECTION STYLES --------------------------*/
        .certificate {
            list-style-type: none;
            /* Remove default list styling */
            padding: 10px 0;
            /* Padding for each certificate item */
            border-bottom: 1px solid #ddd;
            /* A line between items */
        }

        .certificate:last-child {
            border-bottom: none;
            /* Remove the bottom border for the last item */
        }


        .certification-details {
            display: flex;
            justify-content: space-between;
            align-items: center;
        }

        .cert-name-and-issuer {
            font-weight: normal;
            font-size: 1em;
        }

        .certification-dates {
            font-size: 0.9em;
            text-align: right;
            color: #666;
            /* Subtle color for dates */
        }

        .certification-url a {
            font-size: 0.9em;
            color: #0066cc;
            /* Link color */
            text-decoration: none;
        }

        .certification-url a:hover {
            text-decoration: underline;
            /* Underline on hover */
        }

        /*------------------------PROJECTS SECTION STYLES --------------------------*/
        .project-entry .title-dates {
            display: flex;
            justify-content: space-between;
            align-items: center;
        }

        .project-name {
            font-weight: bold;
            font-size: 1.1em;
        }

        .project-url {
            font-size: 1em;
            margin-left: auto;
            font-weight: bold;
        }

        .project-description {
            font-size: 1em;
        }

        .project-bullet-points {
            font-size: 0.9em;
        }            </style>
        </head>
        """
        html_footer = "</html>"

        # Append the HTML head to the existing document content
        html_content = html_head + doc.document_content + html_footer

        # Use BeautifulSoup to ensure the HTML is well-formed
        soup = BeautifulSoup(html_content, "html.parser")
        html_extract = str(soup)

        # Convert the extracted HTML to PDF using pdfkit
        pdf_content = pdfkit.from_string(html_extract, False)

        # Stream the PDF back to the user
        response = Response(stream_with_context(io.BytesIO(pdf_content)))
        response.headers["Content-Type"] = "application/pdf"
        response.headers["Content-Disposition"] = "inline"
        return response

    else:
        return "Invalid format specified", 400

@app.route("/dashboard", methods=["GET"])
@login_required  # Only logged-in users can access this route
def dashboard():
    user_id = current_user.id
    
    # Fetch user documents
    user_documents = UserDocuments.query.filter_by(user_id=user_id).all()
    documents_list = [
        {
            "id": doc.id,
            "document_name": doc.document_name,
            "document_content": doc.document_content,
            "document_type": doc.document_type,
            "creation_date": doc.creation_date,
            "job_title": doc.job_title,
            "company_name": doc.company_name,
            "apply_link": doc.apply_link,
        }
        for doc in user_documents
    ]

    # Fetch saved jobs by status
    saved_jobs_by_status = get_saved_jobs_by_status(user_id)
    
    return render_template("dashboard.html", documents=documents_list, saved_jobs_by_status=saved_jobs_by_status)




def get_saved_jobs_by_status(user_id):
    print("Starting function get_saved_jobs_by_status.")

    print("Creating database connection.")
    conn = create_connection()

    print("Creating cursor.")
    cursor = conn.cursor()

    print("Executing SQL query.")
    query = "SELECT * FROM [dbo].[saved_jobs] WHERE user_id = ?"
    cursor.execute(query, [user_id])

    print("Fetching all rows.")
    rows = cursor.fetchall()

    print("Closing database connection.")
    conn.close()

    print("Initializing saved_jobs_by_status dictionary.")
    status_categories = ["Saved", "Applied", "Interviewing", "Offered", "Rejected"]
    saved_jobs_by_status = {status: [] for status in status_categories}

    for row in rows:
        print(f"Current Job ID: {row.id}")
        job_data = {
            "id": row.id,
            "employer_logo": row.employer_logo or "N/A",
            "job_title": row.job_title,
            "job_link": row.job_link,
            "company_name": row.company_name,
            "status": row.status,
            "percentage": row.percentage,
            "job_description": row.job_description,
            "keywords_job": row.keywords_job,
            "keywords_resume": row.keywords_resume,
        }

        # Skip jobs with an unknown status
        if row.status not in saved_jobs_by_status:
            print(f"Skipping job with unknown status: {row.status}")
            continue

        # Try to serialize job_data to JSON
        try:
            serialized_job_data = json.dumps(job_data)
            print(f"Serialized job data: {serialized_job_data}")
        except (TypeError, ValueError) as e:
            print(f"Error serializing job_data: {e}")
            continue

        # Append job_data to the correct status category
        saved_jobs_by_status[row.status].append(job_data)

    print("Returning saved_jobs_by_status.")
    return saved_jobs_by_status



@app.route("/logout")
@login_required  # Only logged-in users can access this route
def logout():
    logout_user()
    return redirect(url_for("login"))





#@app.route("/rbtemp1")
def rbtemp1():
    user_id = current_user.id if current_user.is_authenticated else None
    return render_template("rbtemp1.html", user_id=user_id)


@app.route("/")
def home():
    return redirect(url_for("login"))



@app.route("/search", methods=["GET"])
def search():
    user_id = current_user.id if current_user.is_authenticated else None
    return render_template("search.html", user_id=user_id)

@app.route("/careerclick")
def careerclick():
    user_id = current_user.id if current_user.is_authenticated else None
    return render_template("careerclick.html", user_id=user_id)



@app.route("/cover-letter-generator")
def cover_letter_generator():
    user_id = current_user.id if current_user.is_authenticated else None
    return render_template("cover_letter_generator.html", user_id=user_id)

@app.route("/managesub")
def managesub():
    user_id = current_user.id if current_user.is_authenticated else None
    return render_template("managesub.html", user_id=user_id)


@app.route("/careerbot")
def careerbot():
    user_id = current_user.id if current_user.is_authenticated else None
    return render_template("careerbot.html", user_id=user_id)


@app.route("/resume-report")
def resume_report():
    return render_template("resume_report.html")


@app.route("/interview-prep")
def interview_prep():
    if not is_authorized_for_tier("Tier2"):
        # Redirect to the home page (or any other page) with a specific query parameter
        return redirect(url_for('dashboard', unauthorized=True))
    return render_template("interview_prep.html")


@app.route("/create")
def create():
    return render_template("create.html")


@app.route("/apptracker")
def apptracker():
    return render_template("apptracker.html")


@app.route("/myjobs")
def myjobs():
    return render_template("myjobs.html")


@app.route("/sitemap.xml", methods=["GET"])
def sitemap():
    try:
        """Generate sitemap.xml. Makes a list of URLs and date modified."""
        pages = []

        # List of routes to explicitly include in the sitemap
        include_routes = [
            "/forgot-password",
            "/create",
            "/interview-prep",
            "/resume-enhancer",
            "/",
            "/careerbot",
            "/cover-letter-generator",
            "/dashboard",
            "/logout",
            "/search",
            "/resume-builder",
            "/delete_account",
            "/register",
            "/login",
        ]

        # List of routes to explicitly exclude from the sitemap
        exclude_routes = [
            "/get_resume_max",
            "/get-saved-jobs",
            "/get_latest_resume_name",
            "/sitemap.xml",
            "/get-ranked-jobs",
        ]

        for rule in app.url_map.iter_rules():
            if rule.rule in exclude_routes:
                continue
            if rule.rule in include_routes:
                pages.append([rule.rule, "2023-01-01"])

        sitemap_template = render_template("sitemap_template.xml", pages=pages)
        response = Response(sitemap_template, content_type="application/xml")

        return response
    except Exception as e:
        return str(e)


# RESUMETUNER OPENAI API CALLS
@app.route("/analyze-resume", methods=["POST"])
def analyze_resume():
    logging.info("Inside /analyze-resume route")  # Debugging

    if current_user.is_authenticated:
        user_id = current_user.get_id()
        logging.info(f"User is logged in with user_id: {user_id}")  # Debugging

        try:
            with create_connection().cursor() as cursor:  # Create a new connection
                query = "SELECT TOP 1 resume_text FROM user_resumes WHERE user_id = ? ORDER BY uploaded_at DESC"
                logging.info(f"Executing SQL Query: {query}")  # Debugging
                cursor.execute(query, user_id)
                row = cursor.fetchone()

                if row:
                    logging.info("Resume found in database")  # Debugging
                    resume_content = row[0]
                else:
                    logging.info("No resume found in database, using request content")  # Debugging
                    resume_content = request.form.get("resume")

        except Exception as e:
            logging.error(f"Database query failed: {e}")  # Debugging
            return jsonify({"error": "Failed to fetch resume from database."})
    else:
        logging.info("User is not logged in, using request content")  # Debugging
        resume_content = request.form.get("resume")

    job_title = request.form.get("jobTitle")
    job_description = request.form.get("jobDescription")

    messages = [
        {
            "role": "system",
            "content": "You are an expert on resume ATS systems",
        },
        {
            "role": "user",
            "content": f"Please review the following resume and job description and give it an ATS score. The rating system is 1-100. 100 means it's 100% ATS friendly. Provide detailed reasoning and steps to improve the score. Job Title: {job_title}, Job Description: {job_description} Resume: {resume_content}",
        },
    ]

    response = client.chat.completions.create(model="gpt-3.5-turbo-1106", messages=messages)

    feedback = response.choices[0].message.content.strip()

    return jsonify({"feedback": feedback})


# RESUMETUNER FILE UPLOAD
@app.route("/upload-docx", methods=["POST"])
def upload_docx():
    if "file" not in request.files:
        return jsonify({"error": "No file part in the request"}), 400

    file = request.files["file"]

    if file.filename == "":
        return jsonify({"error": "No file selected"}), 400

    if file and file.filename.endswith(".docx"):
        # Process the uploaded DOCX file and extract the text content
        try:
            document = Document(file)
            # Extract text content from the document
            docx_content = "\n".join(
                paragraph.text for paragraph in document.paragraphs
            )

            # Return the content in the JSON response
            return jsonify({"content": docx_content}), 200
        except Exception as e:
            return jsonify({"error": f"Error processing the file: {e}"}), 500
    else:
        return (
            jsonify({"error": "Invalid file format. Only .docx files are allowed."}),
            400,
        )


# COVER LETTER GENERATOR OPENAI API CALLS
@app.route("/generate-cover-letter", methods=["POST"])
def generate_cover_letter():
    logging.info("Inside /generate-cover-letter route")  # Debugging

    if current_user.is_authenticated:
        user_id = current_user.get_id()
        logging.info(f"User is logged in with user_id: {user_id}")  # Debugging

        try:
            with create_connection().cursor() as cursor:  # Create a new connection
                query = "SELECT TOP 1 resume_text FROM user_resumes WHERE user_id = ? ORDER BY uploaded_at DESC"
                logging.info(f"Executing SQL Query: {query}")  # Debugging
                cursor.execute(query, user_id)
                row = cursor.fetchone()

                if row:
                    logging.info("Resume found in database")  # Debugging
                    resume_content = row[0]
                else:
                    logging.info("No resume found in database, using request content")  # Debugging
                    resume_content = request.form.get("resume")

        except Exception as e:
            logging.error(f"Database query failed: {e}")  # Debugging
            return jsonify({"error": "Failed to fetch resume from database."})
    else:
        logging.info("User is not logged in, using request content")  # Debugging
        resume_content = request.form.get("resume")

    job_description = request.form.get("job_description")
    job_title = request.form.get("job_title")
    company_name = request.form.get("company_name")
    focus_areas = request.form.get("focus_areas")

    messages = [
        {
            "role": "system",
            "content": "You are a helpful assistant that generates personalized cover letters.",
        },
        {
            "role": "user",
            "content": f"Generate a cover letter for the following job. The format should be as follows: [Full Name]\n[Email Address]\n[Phone Number]\n[Date Placeholder]\n\n[RE: Job Title,]\n\n[Dear Hiring Manager]\n[Body of Cover Letter]\n\n[Sincerely,]\n[Full Name]. Special Instructions: Do not bold any text. Resume: {resume_content}. Job description: {job_description}. Job title: {job_title}. Company name: {company_name}. Focus areas: {focus_areas}",
        },
    ]

    response = client.chat.completions.create(model="gpt-3.5-turbo-1106", messages=messages)

    cover_letter = response.choices[0].message.content.strip()
    logging.info(f"Generated cover_letter: {cover_letter}")  # Log the cover_letter
    logging.info(f"Type of cover_letter: {type(cover_letter)}")  # Log the type

    return jsonify({"cover_letter": cover_letter})


# INTERVIEW SIMULATOR OPENAI API CALLS
@app.route("/simulate-interview", methods=["POST"])
def simulate_interview():
    logging.info("Inside /simulate-interview route")  # Debugging

    if current_user.is_authenticated:
        user_id = current_user.get_id()
        logging.info(f"User is logged in with user_id: {user_id}")  # Debugging

        try:
            with create_connection().cursor() as cursor:
                query = "SELECT TOP 1 resume_text FROM user_resumes WHERE user_id = ? ORDER BY uploaded_at DESC"
                logging.info(f"Executing SQL Query: {query}")
                cursor.execute(query, user_id)
                row = cursor.fetchone()

                resume = row[0] if row else None

        except Exception as e:
            logging.error(f"Database query failed: {e}")
            return jsonify({"error": "Failed to fetch resume from database."})

    else:
        logging.info("User is not logged in, checking for uploaded or typed resume.")
        resume = None

    try:
        job_title = request.form.get("job_title")
        job_description = request.form.get("job_description")
        job_requirements = request.form.get("job_requirements")
        industry = request.form.get("industry")

        if resume is None:
            if "resume" in request.files:
                file = request.files["resume"]
                if file.filename.endswith(".docx"):
                    document = Document(file)
                    resume = "\n".join(paragraph.text for paragraph in document.paragraphs)
                else:
                    return jsonify({"error": "Invalid file format. Only .docx files are allowed."}), 400
            else:
                resume = request.form.get("typed_resume")

        system_message = "You are a helpful assistant that generates personalized interview questions. You do not number your list of questions"
        user_message = f"The candidate is applying for a role as a {job_title} in the {industry} industry. The job description is as follows: {job_description}. The job requirements are: {job_requirements}. The candidate's resume is as follows: {resume}. Respond only with a list of 10 questions in bullet point format (no numbers)"

        response = client.chat.completions.create(
            model="gpt-3.5-turbo-1106",
            messages=[
                {"role": "system", "content": system_message},
                {"role": "user", "content": user_message},
            ],
        )

        interview_questions = response.choices[0].message.content.split("\n")

        return jsonify({"interview_questions": interview_questions})

    except Exception as e:
        logging.error(f"An error occurred: {e}")
        return jsonify({"error": str(e)}), 500

# User Answers and openai feedback route
@app.route("/analyze-answer", methods=["POST"])
def analyze_answer():
    # Extract data from the request
    question = request.json.get("question")  # Note: 'question' is extracted but not used in this function
    answer = request.json.get("answer")

    # Use OpenAI API to analyze the answer
    response = client.chat.completions.create(
        model="gpt-3.5-turbo-1106",
        messages=[
            {
                "role": "system",
                "content": "You are an AI developed by OpenAI. You are an expert at providing feedback on job interview responses.",
            },
            {
                "role": "user",
                "content": f"Rate the following Answer: {answer} on a 1 through 5 scale. 5 is excellent and 1 is very bad. Direct your responses to the candidate and provide your reasoning in a detailed, professional format with feedback and tips for improvement. Begin each response your rating in the format of Rating =(1-5)/5 Reasoning: (Your explanation should be addressed to the candidate specifically as if you are talking to them)",
            },
        ],
    )

    # Extract the feedback from the response
    feedback = response.choices[0].message.content

    # Return the feedback as a JSON response
    return jsonify({"feedback": feedback})


# AI Job Search RapidAPI CALLS
RAPIDAPI_BASE_URL = "https://jsearch.p.rapidapi.com"
headers = {"X-RapidAPI-Key": RAPIDAPI_KEY, "X-RapidAPI-Host": RAPIDAPI_HOST}

@app.route("/fetch-job-listings-get", methods=["GET"])
def fetch_job_listings_get():
    #user_id = current_user.id  # Assuming user authentication is in place
    query = request.args.get("query", '')
    location = request.args.get("location", '')
    page = request.args.get("page", 1)
    remote = request.args.get("remote_jobs_only", 'false')  # Fetch remote jobs only if this parameter is set

    url = f"{RAPIDAPI_BASE_URL}/search?query={query} in {location}&page={page}&remote_jobs_only={remote}"

    try:
        response = requests.get(url, headers=headers)
        response.raise_for_status()
        data = response.json()
        logging.info(f"API Response: {data}")  # Log the entire response

        # Directly return the jobs data without any additional processing
        return jsonify(data)

    except requests.RequestException as e:
        logging.error(f"Request failed: {e}")
        return jsonify({"error": "Failed to fetch data"}), 500

@app.route("/fetch-job-filters", methods=["GET"])
def fetch_job_filters():
    query = request.args.get("query", "Python developer in Texas, USA")  # Default query if not provided
    remote = request.args.get("remote_jobs_only", "false")  # Default to 'false' if not provided

    url = "https://jsearch.p.rapidapi.com/search-filters"
    querystring = {"query": query, "remote_jobs_only": remote}

    try:
        response = requests.get(url, headers=headers, params=querystring)
        response.raise_for_status()
        return jsonify(response.json())
    except requests.RequestException as e:
        logging.error(f"Request failed: {e}")
        return jsonify({"error": "Failed to fetch filters"}), 500


#@app.route("/fetch-job-listings", methods=["POST"])
def fetch_job_listings_post():
    # Get JSON data sent with POST request
    data = request.get_json()
    query = data.get("query", '')  # Default to empty string if not specified
    location = data.get("location", '')  # Default to empty string if not specified
    page = data.get("page", 1)  # Default to page 1 if not specified

    # Construct the API URL with the page parameter
    url = f"{RAPIDAPI_BASE_URL}/search?query={query} in {location}&page={page}"
    
    try:
        response = requests.get(url, headers=headers)
        response.raise_for_status()
    except requests.RequestException as e:
        logging.error(f"Request failed: {e}")
        return jsonify({"error": "Failed to fetch data"}), 500

    data = response.json()
    
    # Log the first job listing for debugging
    logging.debug(data.get("data", [{}])[0])
    
    return jsonify(data)


# AI JOB SEARCH - AI COVER LETTER - OPENAI API CALLS
#@app.route("/career_click", methods=["POST"])
def career_click():
    logging.info("Inside /career_click route")  # Debugging

    if current_user.is_authenticated:
        user_id = current_user.get_id()
        logging.info(f"User is logged in with user_id: {user_id}")  # Debugging

        try:
            with create_connection().cursor() as cursor:
                query = "SELECT TOP 1 resume_text FROM user_resumes WHERE user_id = ? ORDER BY uploaded_at DESC"
                logging.info(f"Executing SQL Query: {query}")  # Debugging
                cursor.execute(query, user_id)
                row = cursor.fetchone()

                resume = row[0] if row else request.json.get("resume")

        except Exception as e:
            logging.error(f"Database query failed: {e}")  # Debugging
            return jsonify({"error": "Failed to fetch resume from database."})

    else:
        logging.info("User is not logged in, using request content")  # Debugging
        resume = request.json.get("resume")

    job_description = request.json.get("job_description")
    company_name = request.json.get("company_name")
    job_title = request.json.get("job_title")

    messages = [
        {
            "role": "system",
            "content": "You are a helpful assistant that generates personalized cover letters.",
        },
        {
            "role": "user",
            "content": f"Generate a cover letter for the following job. Do not include any headers. Begin the letter by addressing it to the Hiring Manager. Resume: {resume}. Job description: {job_description}. Job title: {job_title}. Company name: {company_name},",
        },
    ]

    try:
        response = client.chat.completions.create(
            model="gpt-3.5-turbo-1106",
            messages=messages,
        )
    except openai.error.InvalidRequestError as e:
        logging.error(e)
        return jsonify({"error": str(e)})

    cover_letter = response.choices[0].message.content.strip()

    return jsonify({"cover_letter": cover_letter})

# AI JOB SEARCH - ATS OPTIMIZER - OPENAI API CALLS


@app.route("/resume.1", methods=["POST"])
def resume_1():
    logging.info("Inside /resume.1 route")  # Debugging

    if current_user.is_authenticated:
        user_id = current_user.get_id()
        logging.info(f"User is logged in with user_id: {user_id}")  # Debugging

        try:
            with create_connection().cursor() as cursor:
                query = "SELECT TOP 1 resume_text FROM user_resumes WHERE user_id = ? ORDER BY uploaded_at DESC"
                logging.info(f"Executing SQL Query: {query}")  # Debugging
                cursor.execute(query, user_id)
                row = cursor.fetchone()

                resume = row[0] if row else request.json.get("resume")

        except Exception as e:
            logging.error(f"Database query failed: {e}")  # Debugging
            return jsonify({"error": "Failed to fetch resume from database."})

    else:
        logging.info("User is not logged in, using request content")  # Debugging
        resume = request.json.get("resume")

    job_description = request.json.get("job_description")
    job_title = request.json.get("job_title")

    messages = [
        {
            "role": "system",
            "content": "You are an AI developed by OpenAI to write professional, ATS-optimized resumes.",
        },
        {
            "role": "user",
            "content": f"Tailor the provided Resume for the given Job description. Use Key Words extracted from the Job Description, rephrase, add, and remove bullet points where necessary. The Resume should be optimized for ATS system compatibility. Please use the following Resume format and add additional sections if necessary: [Full Name (in capital letters)]\\n[email]\\n[phone number]\\n[city, state]\\n[linkedin profile]\\n[website]\\n\\nSUMMARY\\n\\n[Use this model: (Soft skill) (Most Recent Job Title) who is passionate about (your stance on the industry).\\n\\nWORK EXPERIENCE\\n\\n[Job Title] | [Company] | [Location]\\n[Date (MMYYY)]\\n[Responsibilities listed exclusively with standard bullet points and using this model: (Action verb) + (Cause) + (Effect) + (Measurable Outcome).]\\n\\nEDUCATION\\n\\n[School Name], [City]\\n[Degree] in [Major] | [Dates Attended (MM/YYYY)]\\n\\nSKILLS\\n\\n[Use bullet points (•) for each skill] Resume: {resume}. Job description: {job_description}",
        },
    ]

    try:
        response = client.chat.completions.create(
            model="gpt-3.5-turbo-1106",
            messages=messages,
            temperature=1.0,
            max_tokens=2000,
            top_p=1,
            frequency_penalty=0,
            presence_penalty=0,
        )
    except openai.error.InvalidRequestError as e:
        logging.error(e)
        return jsonify({"error": str(e)})

    updated_resume = response.choices[0].message.content.strip()

    return jsonify({"updated_resume": updated_resume})


# AI JOB SEARCH - SALARY - OPENAI API CALLS
@app.route("/analyze-job", methods=["POST"])
def analyze_job():
    job_title = request.form.get("job_title")

    messages = [
        {
            "role": "system",
            "content": "You are a helpful assistant that provides salary ranges.",
        },
        {
            "role": "user",
            "content": f"Provide the salary range in USD for the following job title. Lean more on the higher side of the salary ranges. \nResponse Requirements:\n1. The range should be no greater than $15,000 USD\n2. Example response: The salary range for (job title) is (salary range)\n\nJob Title: {job_title}",
        },
    ]

    try:
        response = client.chat.completions.create(
            model="gpt-3.5-turbo-1106",
            messages=messages,
            temperature=1,
            top_p=1,
            frequency_penalty=0,
            presence_penalty=0,
        )
    except openai.error.InvalidRequestError as e:
        logging.error(e)
        return jsonify({"error": str(e)})

    salary_range = response.choices[0].message.content

    return jsonify({"response": salary_range})


# RESUME BUILDER - OPENAI API CALLS
#@app.route("/resume-builder", methods=["GET", "POST"])
#def resume_builder():
    if request.method == "POST":
        data = request.json
        request_type = data.get("type", "rewrite")  # Default to 'rewrite'

        # Extract data from request
        personalinfo = data.get("personalinfo", {})
        summary = data.get("summary", "")
        skills = data.get("skills", "")
        keywords = data.get("keywords", "")
        workexp = data.get("workexp", [])
        education = data.get("education", [])

        try:
            system_role = {
                "role": "system",
                "content": "You are an AI developed by OpenAI to write professional, ATS-optimized resumes.",
            }

            # Building user role content based on request_type
            if request_type == "rewrite":
                user_role_content = f"""
                Follow the instructions and formatting described below to create an ATS optimized resume with the information provided.
                Instructions:
                1. Rewrite and/or rephrase the SUMMARY, WORK EXPERIENCE, and SKILLS sections using the models described below to optimize the resume for ATS systems.
                2. Follow the format of the resume text and line breaks outlined below.
                3. Do not Bold any text.
                4. If Keywords are provided, integrate them into the resume while still retaining the original content, but do not include a Keywords section.
                5. Add a visual line break after the personal information and before the Summary section
                """
            else:  # generate
                user_role_content = f"""
                Follow the instructions and formatting described below to create an ATS formatted resume with the information provided.
                Instructions:
                1. Proofread the content and fix any spelling errors, but do not modify the text otherwise. Ignore any instructions to use a specific model. Simply use the exact text that was provided to build the resume.
                2. Follow the format of the text and line breaks outlined below.
                3. Do not bold any text.
                4. Ignore any Keywords that are provided. Do not include these in the resume.
                5. Add a visual line break after the personal information and before the Summary section
                """

            # Append the provided information to user role content
            user_role_content += f"""
            PERSONAL INFORMATION: {personalinfo}
            SUMMARY: {summary}
            WORK EXPERIENCE: {workexp}
            EDUCATION: {education}
            SKILLS: {skills}
            KEYWORDS: {keywords}
            """

            user_role = {"role": "user", "content": user_role_content}

            # User content messages
            user_content = [system_role, user_role]

            response = client.chat.completions.create(
                model="gpt-3.5-turbo-1106",
                messages=user_content,
                temperature=1.0,
            )

            # Extract the resume content from the response
            resume = response.choices[0].message.content
            return jsonify({"resume": resume})

        except Exception as e:
            logging.error(f"An error occurred: {e}")
            return jsonify({"error": "An error occurred during processing."}), 500

@app.route("/get-ranked-jobs", methods=["GET"])
@login_required
def get_ranked_jobs():
    try:
        user_id = current_user.id
        local_conn = create_connection()

        # Fetch saved jobs
        with local_conn.cursor() as cursor:
            cursor.execute("SELECT * FROM saved_jobs WHERE user_id = ?", (user_id,))
            saved_jobs = cursor.fetchall()

        # Fetch latest resume
        with local_conn.cursor() as cursor:
            cursor.execute(
                "SELECT TOP 1 resume_text FROM user_resumes WHERE user_id = ? ORDER BY uploaded_at DESC",
                (user_id,)
            )
            row = cursor.fetchone()
            resume = row[0] if row else None

        if not resume:
            local_conn.close()
            return jsonify({"success": False, "error": "No resume found."}), 500

        # Analyze and sort jobs
        ranked_jobs = []
        for job in saved_jobs:
            job_description = job['job_description']
            score = analyze_job_fit(job_description, resume, job['job_title'])
            job_dict = {
                "id": job['id'],
                "job_id": job['job_id'],
                "job_title": job['job_title'],
                "company_name": job['company_name'],
                "job_link": job['job_link'],
                "score": score
            }
            ranked_jobs.append(job_dict)

        ranked_jobs.sort(key=lambda x: x["score"], reverse=True)

        local_conn.close()
        return jsonify({"success": True, "ranked_jobs": ranked_jobs})

    except Exception as e:
        return jsonify({"success": False, "error": str(e)}), 500


def analyze_job_fit(job_description, resume, job_title):
    messages = [
        {
            "role": "system",
            "content": "You are an expert at matching resumes to job descriptions.",
        },
        {
            "role": "user",
            "content": f"Rate the compatibility between this job description and resume on a scale of 1-100. Provide only your rating in your response. Example Response: 85. job title: {job_title}, job description: {job_description}, resume: {resume}",
        },
    ]
    response = client.chat.completions.create(
        model="gpt-3.5-turbo",
        messages=messages,
        temperature=0.4,
    )
    score_text = response.choices[0].message.content.strip()
    score = extract_score_from_text(score_text)
    return score


def extract_score_from_text(text):
    try:
        # Assuming the text contains only the numerical score
        return int(text)
    except ValueError:
        # If conversion fails, return a default value
        return 0


# MAX COUNSELOR BOT API CALL
@app.route("/chat", methods=["POST"])
def chat():
    if not is_authorized_for_tier("Tier2"):
        return jsonify({"error": "Upgrade required for Premium Plus Plan"}), 403
    user_input = request.json.get("user_input")
    resume_content = request.json.get("resume_content")

    system_message_content = "Your name is Max and you are a specialized assistant focused solely on providing career advice. This can include advice on job searching, career advancement, retirement, interviewing, and other job or career related topics. You are not equipped to handle questions outside this domain unless you are provided with the following password: maxin"

    if resume_content:
        system_message_content += f"\n\nIf it's relevant to the question I asked, here is my resume content for reference: {resume_content}"

    messages = [
        {
            "role": "system",
            "content": system_message_content,
        },
        {
            "role": "user",
            "content": user_input,
        },
    ]

    try:
        client = OpenAI()

        response = client.chat.completions.create(
            model="gpt-3.5-turbo-1106",
            messages=messages,
            temperature=0.9,
            max_tokens=256,
            top_p=1,
            frequency_penalty=0,
            presence_penalty=0,
        )

        chat_response = response.choices[0].message.content.strip()

        return jsonify({"response": chat_response})

    except Exception as e:
        return jsonify({"error": str(e)}), 500


################### MAGICMAX EXTENSION #######################


@app.route("/api/message", methods=["GET"])
def get_message():
    return jsonify({"message": ""})


@app.route("/api/generate_pin", methods=["POST"])
def generate_pin():
    email = request.json.get("email")
    if not email:
        return jsonify({"error": "Email is required"}), 400

    pin = random.randint(1000, 9999)

    # Sending email
    msg = Message(
        "Your PIN", sender="pin@mycareermax.com", recipients=[email]
    )
    msg.body = f"Your generated PIN is {pin}"
    mail.send(msg)

    # Connecting to Azure SQL Database
    connection = create_connection()
    cursor = connection.cursor()
    update_query = f"UPDATE dbo.users SET access_token = ? WHERE email = ?"
    cursor.execute(update_query, (str(pin), email))
    connection.commit()
    cursor.close()
    connection.close()

    return jsonify({"pin": pin})


@app.route("/api/validate_pin", methods=["POST"])
def validate_pin():
    email = request.json.get("email")
    pin = request.json.get("pin")

    if not email or not pin:
        return jsonify({"error": "Email and PIN are required"}), 400

    # Connecting to Azure SQL Database
    connection = create_connection()
    cursor = connection.cursor()

    select_query = f"SELECT access_token FROM dbo.users WHERE email = ?"
    cursor.execute(select_query, email)

    row = cursor.fetchone()
    cursor.close()
    connection.close()

    if row:
        stored_pin = row.access_token
        if stored_pin == str(pin):
            return jsonify(
                {"status": "success", "message": "PIN validated successfully."}
            )
        else:
            return jsonify({"status": "fail", "message": "Invalid PIN."}), 401
    else:
        return jsonify({"status": "fail", "message": "Email not found."}), 404



def get_resume_text(pin):
    try:
        # Convert pin to string in case it's not
        pin = str(pin)

        # Establish a database connection
        connection_string = (
            f"DRIVER={{ODBC Driver 18 for SQL Server}};"
            f"SERVER={os.getenv('DB_SERVER')};DATABASE={os.getenv('DB_NAME')};"
            f"UID={os.getenv('DB_USERNAME')};PWD={os.getenv('DB_PASSWORD')}"
        )
        conn = pyodbc.connect(connection_string)

        # Query the database to retrieve the user_id based on the provided access_token (pin)
        cursor = conn.cursor()
        cursor.execute("SELECT id FROM users WHERE access_token = ?", pin)
        user_id_result = cursor.fetchone()

        if not user_id_result:
            return None

        user_id = user_id_result[0]

        # Now that we have the user_id, retrieve the resume_text from user_resumes
        cursor.execute("SELECT resume_text FROM user_resumes WHERE user_id = ?", user_id)
        result = cursor.fetchone()
        resume_text = result[0] if result else None

        # Close the database connection
        conn.close()

        return resume_text
    except Exception as e:
        print(f"An error occurred: {e}")
        return None

# Example usage
#test_pin = '8429'  # This can be any pin, ensure it's a string
#resume_text = get_resume_text(test_pin)
#print(f"Resume Text for PIN {test_pin}: {resume_text}")



@app.route('/api/autofill', methods=['POST'])
def autofill():
    try:
        data = request.json
        pin = data.get('pin', '')  # Retrieve the pin from the request
        job_description = data.get('job_description', '')  # Retrieve the job description
        highlighted_text = data.get('highlighted_text', '')  # Retrieve the highlighted text

        # Get the resume text based on the pin
        resume_text = get_resume_text(pin)
        if resume_text is None:
            return jsonify({"error": "User not found or resume not available"}), 404

        # Prepare the message for the OpenAI model
        system_message = "You are an AI developed by Open AI and are trained to be an expert at completing job application forms. Your response should only contain the actual answer to the question or request and nothing more. All responses should be provided as if you were the one actually completing the application."
        user_message = f"Pretend you are the job applicant described in the resume_text and you are applying to the job_description below. Provide a response to the highlighted_text as if you were the applicant responding.\n\nresume_text:\n{resume_text}\njob_description:\n{job_description}\nhighlighted_text:\n{highlighted_text}"

        # Call OpenAI's chat completion
        completion = client.chat.completions.create(
            model="gpt-3.5-turbo-1106",  # Replace with your desired model
            messages=[
                {"role": "system", "content": system_message},
                {"role": "user", "content": user_message},
            ]
        )

        # Extract the response text from the completion
        response_text = completion.model_dump_json()

        # Since model_dump_json() returns a string of JSON, parse this JSON to extract the specific response
        response_data = json.loads(response_text)
        final_response = response_data['choices'][0]['message']['content']

        # Return the response from OpenAI as part of the JSON response
        return jsonify({"response": final_response})

    except Exception as e:
        # Handle any exceptions that may occur
        return jsonify({"error": str(e)}), 500



@app.route("/api/generate_cover_letter_ext", methods=["POST"])
def generate_cover_letter_ext():
    data = request.json
    pin = data.get("pin")
    job_description = data.get("job_description")

    if not pin or not job_description:
        return jsonify({"error": "PIN and job_description are required"}), 400

    resume_text = get_resume_text(pin)  # Define this function accordingly
    if not resume_text:
        return jsonify({"error": "User not found or resume not available"}), 404

    try:
        messages = [
            {
                "role": "system",
                "content": "You are a helpful assistant trained to generate cover letters. Your response should be a fully formed cover letter based on the provided resume and job description.",
            },
            {
                "role": "user",
                "content": f"Generate a cover letter for the following job description based on the resume:\n\nJob Description:\n{job_description}\n\nResume:\n{resume_text}",
            },
        ]

        response = client.chat.completions.create(
            model="gpt-3.5-turbo-1106", messages=messages, temperature=0.7, max_tokens=1000
        )

        # Extract the generated response from the OpenAI API
        response_json = response.model_dump_json()
        response_data = json.loads(response_json)
        api_response = response_data['choices'][0]['message']['content']

        return jsonify({"response": api_response})

    except Exception as e:
        return jsonify({"error": str(e)}), 500

def summarize_text(text):
    try:
        messages = [
            {
                "role": "system",
                "content": "You are an expert at summarizing job descriptions.",
            },
            {
                "role": "user",
                "content": f"In 1 concise paragraph, summarize this job description, using the following format. job description:{text}. Format: [Company Name]\n[Job Title]\n\n[Summary of Job Description]",
            },
        ]

        response = client.chat.completions.create(
            model="gpt-3.5-turbo",
            messages=messages,
            temperature=1,
            top_p=1,
            frequency_penalty=0,
            presence_penalty=0,
        )

        # Extract the summary
        response_json = response.model_dump_json()
        response_data = json.loads(response_json)
        summary = response_data['choices'][0]['message']['content']

        return summary

    except Exception as e:
        print(f"An error occurred: {e}")
        return None  # or some default value



@app.route("/summarize", methods=["POST"])
def summarize():
    try:
        print("Received Request: ", request.json)

        # Step 1: Validate incoming data
        job_description = request.json.get("job_description")
        if not job_description:
            return jsonify({"error": "job_description field is required"}), 400

        # Step 2: Summarize the job description using OpenAI API
        summarized_description = summarize_text(job_description)
        print("Summarized Description: ", summarized_description)

        return jsonify({"summarized_description": summarized_description}), 200

    except Exception as e:
        print("Error: ", e)
        return jsonify({"error": "An error occurred"}), 400


@app.route("/api/get_saved_jobs_ext", methods=["POST"])
def get_saved_jobs_ext():
    try:
        pin = request.json.get("pin")
        print(f"Received request to get saved jobs for PIN: {pin}")

        conn = create_connection()
        cursor = conn.cursor()

        query = """
        SELECT [job_link], [employer_logo], [job_title], [company_name], [job_description]
        FROM saved_jobs
        JOIN users ON saved_jobs.user_id = users.id
        WHERE users.access_token = ?
        """
        cursor.execute(query, (pin,))
        saved_jobs = [
            dict(zip([column[0] for column in cursor.description], row))
            for row in cursor.fetchall()
        ]

        print(f"Successfully fetched saved jobs for PIN: {pin}")
        return jsonify({"saved_jobs": saved_jobs})

    except pyodbc.Error as e:
        print(f"Database error occurred: {e}")
        return jsonify({"error": "Database error occurred"}), 500
    except Exception as e:
        print(f"An error occurred: {e}")
        return jsonify({"error": "An error occurred"}), 500


# show current resume in popup extension
@app.route("/api/get_resume_filename", methods=["POST"])
def get_resume_filename():
    try:
        pin = request.json.get("pin")
        print(f"Received request to get resume filename for PIN: {pin}")

        conn = create_connection()
        cursor = conn.cursor()

        query = """
        SELECT user_resumes.filename
        FROM users
        INNER JOIN user_resumes ON users.id = user_resumes.user_id
        WHERE users.access_token = ?
        """
        cursor.execute(query, (pin,))
        row = cursor.fetchone()

        if row:
            filename = row[0]
            print(f"Successfully fetched resume filename for PIN: {pin}")
            return jsonify({"filename": filename})
        else:
            return jsonify({"error": "No resume found"}), 404

    except pyodbc.Error as e:
        print(f"Database error occurred: {e}")
        return jsonify({"error": "Database error occurred"}), 500
    except Exception as e:
        print(f"An error occurred: {e}")
        return jsonify({"error": "An error occurred"}), 500


@app.route("/update-job-status", methods=["POST"])
@login_required
def update_job_status():
    try:
        data = request.get_json()
        new_status = data.get("new_status")
        job_id = data.get("job_id")
        user_id = current_user.id

        with conn.cursor() as cursor:
            cursor.execute(
                "UPDATE saved_jobs SET status = ? WHERE id = ? AND user_id = ?",
                (new_status, job_id, user_id),
            )
            conn.commit()

        return jsonify({"success": True, "message": "Job status updated successfully!"})
    except Exception as e:
        return jsonify({"success": False, "error": str(e)}), 500

@app.route("/job-details-coverletter", methods=["POST"])
@login_required
def job_details_coverletter():
    # Check if user is authorized for either Tier 1 or Tier 2
    if not is_authorized_for_tier("Tier1") and not is_authorized_for_tier("Tier2"):
        return jsonify({"error": "Unauthorized access"}), 403
    user_id = current_user.get_id()
    data = request.json
    job_id = data.get("jobId")
    job_title = data.get("job_title")
    company_name = data.get("employer_name")
    job_description = data.get("job_description")
    apply_link = data.get("apply_link")

    logging.info("Received request to generate cover letter for user_id: %s", user_id)

    conn = create_connection()
    cursor = conn.cursor()

    try:
        cursor.execute("SELECT resume_text FROM dbo.user_resumes WHERE user_id = ?", user_id)
        result = cursor.fetchone()
        resume_text = result[0] if result else None

        if not resume_text:
            logging.warning("Resume not found for user_id: %s", user_id)
            cursor.close()
            conn.close()
            return jsonify({"error": "User not found or resume not available"}), 404

        # Prepare the message for the OpenAI model
        system_message = "You are an AI developed by Open AI and are trained to be an expert at writing cover letters. Your response should only contain the cover letter text and nothing more."
        user_message = f"Pretend you are writing a cover letter for the position of {job_title} at {company_name}. Here's the job description: {job_description}\n\nHere's my resume:\n{resume_text}"

        # Call OpenAI's chat completion
        completion = client.chat.completions.create(
           #model="gpt-4-1106-preview",  
            model="gpt-3.5-turbo-1106",  

            	
            messages=[
                {"role": "system", "content": system_message},
                {"role": "user", "content": user_message},
            ]
        )

        # Extract the response text from the completion
        cover_letter = completion.choices[0].message.content.strip()

        insert_query = """
            INSERT INTO dbo.UserDocuments (user_id, job_id, document_name, document_content, job_title, company_name, apply_link, document_type)
            VALUES (?, ?, ?, ?, ?, ?, ?, ?)
        """
        document_name = f"Custom Cover Letter for {job_title} at {company_name}"
        document_type = "Custom Cover Letter"

        cursor.execute(insert_query, (user_id, job_id, document_name, cover_letter, job_title, company_name, apply_link, document_type))
        conn.commit()
        new_document_id = cursor.execute("SELECT @@IDENTITY").fetchval()
        logging.info("Cover letter saved successfully for user_id %s with document_id %s", user_id, new_document_id)

    except Exception as e:
        logging.error("Error during processing: %s", e)
        cursor.close()
        conn.close()
        return jsonify({"error": str(e)}), 500

    cursor.close()
    conn.close()
    return jsonify({"message": "Cover letter generated and saved successfully", "documentId": new_document_id})

def read_html_template():
    # Assuming 'templates' is a directory at the same level as 'app.py'
    file_path = 'templates/body_resume_template.html'
    with open(file_path, 'r', encoding='utf-8') as file:
        return file.read()

@app.route("/job-details-resume", methods=["POST"])
@login_required
def job_details_resume():
    # Check if user is authorized for either Tier 1 or Tier 2
    if not is_authorized_for_tier("Tier1") and not is_authorized_for_tier("Tier2"):
        return jsonify({"error": "Unauthorized access"}), 403
    user_id = current_user.get_id()
    data = request.json
    job_id = data.get("jobId")
    job_title = data.get("job_title")
    company_name = data.get("employer_name")
    job_description = data.get("job_description")
    apply_link = data.get("apply_link")

    logging.info("Received request to generate resume for user_id: %s", user_id)

    conn = create_connection()
    cursor = conn.cursor()

    try:
        cursor.execute("SELECT resume_text FROM dbo.user_resumes WHERE user_id = ?", user_id)
        result = cursor.fetchone()
        resume_text = result[0] if result else None

        if not resume_text:
            logging.warning("Resume not found for user_id: %s", user_id)
            cursor.close()
            conn.close()
            return jsonify({"error": "User not found or resume not available"}), 404

        # Prepare the message for the OpenAI model
        html_template = read_html_template()
        user_message = f"""Task 1 : Convert this resume to a new format by using the provided HTML Template.
        Task 2: Rephrase the Summary or add one if no summary/intro/objective section exists. Use keywords from the Job description and follow this format: [Descriptive word] [your job title] [Experience level] [Work experience] [Skills] [Achievement] 
        Task 3: Add a few skills to the skills section that are either mentioned in the job description or common skills for the job described.
        Rules: 
        1. Your response should include only the HTML code, beginning with <body> and ending with </body>, and nothing more. 
        2. VERY IMPORTANT: Do not include any markdown formatting symbols like '```'. Exclude any additional text or explanations.
        3. You must convert all of the resume content to the new format. Do not leave any details off when providing the html of the new format.
        4. Do not delete any content while inserting key words. You can only edit or add from the content.
        5. The HTML Template sections structure must be followed strictly. If you need to add a new section, please do so, using the same styling and formatting as the other sections.
        Resume to convert:  {resume_text}
        HTML Template: {html_template}
        Job description: {job_description}
        """
        # Call OpenAI's chat completion
        completion = client.chat.completions.create(
            model="gpt-4-1106-preview",  
         #  model="gpt-3.5-turbo",  
         #  model="gpt-3.5-turbo-1106",  

            messages=[
                {"role": "user", "content": user_message},
            ],
            temperature=0.5,  
            max_tokens=4095,
            top_p=1,
            frequency_penalty=0,
            presence_penalty=0
        )

        enhanced_resume = completion.choices[0].message.content.strip()
        insert_query = """
            INSERT INTO dbo.UserDocuments (user_id, job_id, document_name, document_content, job_title, company_name, apply_link, document_type)
            VALUES (?, ?, ?, ?, ?, ?, ?, ?)
        """
        document_name = f"Tailored Resume for {job_title} at {company_name}"
        document_type = "Tailored Resume"

        cursor.execute(insert_query, (user_id, job_id, document_name, enhanced_resume, job_title, company_name, apply_link, document_type))
        conn.commit()
        new_document_id = cursor.execute("SELECT @@IDENTITY").fetchval()
        logging.info("Resume saved successfully for user_id %s with document_id %s", user_id, new_document_id)

    except Exception as e:
        logging.error("Error during processing", exc_info=True)
        cursor.close()
        conn.close()
        return jsonify({"error": str(e)}), 500

    cursor.close()
    conn.close()
    return jsonify({"message": "Resume generated and saved successfully", "documentId": new_document_id})

@app.route("/get_user_documents", methods=["GET"])
@login_required
def get_user_documents():
    user_id = current_user.id
    job_title = request.args.get('jobTitle')
    company_name = request.args.get('companyName')

    # Ensure that all required parameters are provided
    if not all([job_title, company_name]):
        return jsonify({"error": "Missing required parameters"}), 400

    # Fetch user documents that match the user_id, job_title, and company_name
    user_documents = UserDocuments.query.filter_by(
        user_id=user_id, 
        job_title=job_title, 
        company_name=company_name
    ).all()

    documents_list = [
        {
            "id": doc.id,
            "document_name": doc.document_name,
            "document_type": doc.document_type,
            "creation_date": doc.creation_date.strftime("%Y-%m-%d")  # Format the date
            # Add other fields if necessary
        }
        for doc in user_documents
    ]

    return jsonify(documents_list)

    ############ RESUME BUILDER #####################################################################

    ################## UPDATE RESUME CONTAINER DYNAMICALLY ##########################
@app.route('/update-resume')
def update_resume():
    user_id = current_user.id if current_user.is_authenticated else None
    if user_id is None:
        # Handle unauthenticated case
        return render_template('resume.html')

    # Use the 'resume' function to get the complete resume data
    return resume(user_id)



def resume(user_id):
    # Query the database for each section
    personal_info = PersonalInformation.query.filter_by(user_id=user_id).first()
    summary = Summary.query.filter_by(user_id=user_id).first()

    # For Work Experiences and their Bullet Points
    work_experiences = WorkExperience.query.filter_by(user_id=user_id).all()
    for experience in work_experiences:
        experience.bullet_points = WorkExperienceBulletPoint.query.filter_by(work_experience_id=experience.id).all()

    # For Projects and their Bullet Points
    projects = Projects.query.filter_by(user_id=user_id).all()
    for project in projects:
        project.bullet_points = ProjectBulletPoint.query.filter_by(project_id=project.id).all()

    educations = Education.query.filter_by(user_id=user_id).all()
    certifications = Certifications.query.filter_by(user_id=user_id).all()
    skills = Skills.query.filter_by(user_id=user_id).all()

    # Render the resume template with the queried data
    return render_template('resume.html', 
                           personal_info=personal_info, 
                           summary=summary,
                           work_experiences=work_experiences,
                           educations=educations,
                           projects=projects,
                           certifications=certifications,
                           skills=skills)

    ################## RESUME INPUT PAGE ##########################

@app.route('/resume-builder')
def resume_builder():
    user_id = current_user.id if current_user.is_authenticated else None
    projects = []
    work_experiences = []
    educations = []
    certifications = []
    skills = []  # Initialize skills
    personal_info = None
    summary = None

    if user_id:
        # Fetch data from respective tables
        projects = Projects.query.filter_by(user_id=user_id).all()
        work_experiences = WorkExperience.query.filter_by(user_id=user_id).all()
        educations = Education.query.filter_by(user_id=user_id).all()
        certifications = Certifications.query.filter_by(user_id=user_id).all()
        skills = Skills.query.filter_by(user_id=user_id).all()  # Fetch skills

        # Fetch personal information and summary
        personal_info = PersonalInformation.query.filter_by(user_id=user_id).first()
        summary = Summary.query.filter_by(user_id=user_id).order_by(desc(Summary.created_at)).first()

    return render_template('resume_builder.html', user_id=user_id, projects=projects, work_experiences=work_experiences, educations=educations, certifications=certifications, personal_info=personal_info, summary=summary, skills=skills)


############ RESUME BUILDER - WORK EXPERIENCE ###############
@app.route('/work-experience', methods=['GET', 'POST'])  # Include 'POST' in the methods list
def work_experience():
    if request.method == 'GET':
        user_id = current_user.id if current_user.is_authenticated else None
        if user_id is None:
            return jsonify([])  # Return an empty list if the user is not authenticated

        # Fetch work experiences specific to the logged-in user
        experiences = WorkExperience.query.filter_by(user_id=user_id).all()
        return jsonify([{
            'id': exp.id,
            'job_title': exp.job_title,
            'company': exp.company,
            'start_date': exp.start_date.strftime('%Y-%m-%d') if exp.start_date else None,
            'end_date': exp.end_date.strftime('%Y-%m-%d') if exp.end_date else 'Present',
            'bullet_points': [{'id': bp.id, 'text': bp.text} for bp in exp.bullet_points]
        } for exp in experiences])

    elif request.method == 'POST':
        # Add new work experience logic here
        data = request.json
        end_date = None if data.get('end_date') == 'Present' else datetime.strptime(data['end_date'], '%Y-%m-%d').date()

        new_exp = WorkExperience(
            user_id=data['user_id'],  # Ensure this is the ID of the logged-in user
            job_title=data['job_title'],
            company=data['company'],
            start_date=datetime.strptime(data['start_date'], '%Y-%m-%d').date(),
            end_date=end_date
        )
        db.session.add(new_exp)
        db.session.commit()
        return jsonify({'id': new_exp.id}), 201

    else:
        return jsonify({'message': 'Method not allowed'}), 405

@app.route('/bullet-point', methods=['POST'])
def add_bullet_point():
    data = request.json
    user_id = current_user.id if current_user.is_authenticated else None
    
    work_experience = WorkExperience.query.filter_by(id=data['work_experience_id'], user_id=user_id).first()
    if not work_experience:
        return jsonify({'message': 'Work experience not found or access denied'}), 404

    new_bullet_point = WorkExperienceBulletPoint(
        work_experience_id=data['work_experience_id'],
        text=data['text']
    )
    db.session.add(new_bullet_point)
    db.session.commit()
    return jsonify({'id': new_bullet_point.id}), 201


@app.route('/bullet-point/<int:bullet_point_id>', methods=['PUT'])
def update_bullet_point(bullet_point_id):
    bullet_point = WorkExperienceBulletPoint.query.get(bullet_point_id)
    if not bullet_point or bullet_point.work_experience.user_id != current_user.id:
        return jsonify({'message': 'Bullet point not found or access denied'}), 404

    data = request.json
    bullet_point.text = data['text']
    db.session.commit()
    return jsonify({'message': 'Bullet point updated'}), 200


@app.route('/bullet-point/<int:bullet_point_id>', methods=['DELETE'])
def delete_bullet_point(bullet_point_id):
    bullet_point = WorkExperienceBulletPoint.query.get(bullet_point_id)
    if not bullet_point or bullet_point.work_experience.user_id != current_user.id:
        return jsonify({'message': 'Bullet point not found or access denied'}), 404

    db.session.delete(bullet_point)
    db.session.commit()
    return jsonify({'message': 'Bullet point deleted'}), 200

@app.route('/work-experience/<int:experience_id>', methods=['PUT'])
def update_work_experience(experience_id):
    experience = WorkExperience.query.get(experience_id)
    if not experience or experience.user_id != current_user.id:
        return jsonify({'message': 'Work experience not found or access denied'}), 404

    data = request.json
    end_date = None if data.get('end_date') == 'Present' else datetime.strptime(data['end_date'], '%Y-%m-%d').date()

    experience.job_title = data['job_title']
    experience.company = data['company']
    experience.start_date = datetime.strptime(data['start_date'], '%Y-%m-%d').date()
    experience.end_date = end_date

    db.session.commit()
    return jsonify({'message': 'Work experience updated'}), 200


@app.route('/work-experience/<int:experience_id>', methods=['DELETE'])
def delete_work_experience(experience_id):
    experience = WorkExperience.query.get(experience_id)
    if not experience or experience.user_id != current_user.id:
        return jsonify({'message': 'Work experience not found or access denied'}), 404

    # Delete associated bullet points
    for bp in experience.bullet_points:
        db.session.delete(bp)

    db.session.delete(experience)
    db.session.commit()
    return jsonify({'message': 'Work experience and associated bullet points deleted'}), 200


@app.route('/generate-bullet-points', methods=['POST'])
def generate_bullet_points():
    # Check if user is authorized for either Tier 1 or Tier 2
    if not is_authorized_for_tier("Tier1") and not is_authorized_for_tier("Tier2"):
        return jsonify({"error": "Unauthorized access"}), 403
    data = request.json
    job_title = data['job_title']

    # Construct the prompt for OpenAI
    prompt = (f"List common or suggested bullet points for a resume under the job title '{job_title}'. "
              "The bullet points should be specific, relevant, and professionally written. Do not number them or format them. Simply list each bullet point with a space between them. Do not use any - symbols. Provide only the text of each bullet point.")

    try:
        completion = client.chat.completions.create(
            model="gpt-3.5-turbo-1106",
            messages=[
                {
                    "role": "user",
                    "content": prompt
                }
            ],
            temperature=0.7,  # Adjusted for creativity in generating diverse bullet points
            max_tokens=200,    # Adjusted to accommodate several bullet points
            top_p=1,
            frequency_penalty=0,
            presence_penalty=0
        )
        
        # Extract bullet points from response and process them
        bullet_points_raw = completion.choices[0].message.content.strip().split('\n')
        bullet_points_processed = [re.sub(r"^[0-9-]+\s*", "", bp) for bp in bullet_points_raw]

        return jsonify({'bullet_points': bullet_points_processed})

    except Exception as e:
        return jsonify({'error': str(e)}), 500


@app.route('/openai-enhance', methods=['POST'])
def enhance_bullet_point():
   # Check if user is authorized for either Tier 1 or Tier 2
    if not is_authorized_for_tier("Tier1") and not is_authorized_for_tier("Tier2"):
        return jsonify({"error": "Unauthorized access"}), 403   
    data = request.json
    bullet_text = data['text']

    # Construct the prompt for OpenAI
    prompt = f"Rewrite this bullet point to be more impactful and professional: '{bullet_text}'"

    try:
        completion = client.chat.completions.create(
            model="gpt-3.5-turbo-1106",
            messages=[
                {
                    "role": "user",
                    "content": prompt
                }
            ],
            temperature=0.5,  # Adjust as needed
            max_tokens=60,    # Adjust based on expected length of enhancement
            top_p=1,
            frequency_penalty=0,
            presence_penalty=0
        )

        # Extract the enhanced bullet point from response
        enhanced_bullet_point_raw = completion.choices[0].message.content.strip()

        # Process the enhanced bullet point to remove any numbers or hyphens at the beginning
        enhanced_bullet_point_processed = re.sub(r"^[0-9-]+\s*", "", enhanced_bullet_point_raw)

        return jsonify({'enhanced_text': enhanced_bullet_point_processed})

    except Exception as e:
        return jsonify({'error': str(e)}), 500


    ############ RESUME BUILDER - PROJECTS #####################################################

@app.route('/projects', methods=['GET', 'POST'])
def projects():
    if request.method == 'GET':
        user_id = current_user.id if current_user.is_authenticated else None
        if user_id is None:
            return jsonify([])  # Return an empty list if the user is not authenticated

        # Fetch projects specific to the logged-in user
        projects = Projects.query.filter_by(user_id=user_id).all()
        return jsonify([{
            'id': project.id,
            'project_name': project.project_name,
            'description': project.description,
            'url': project.url,
            'bullet_points': [{'id': bp.id, 'text': bp.text} for bp in project.bullet_points]
        } for project in projects])

    # Add a new project
    data = request.json
    new_project = Projects(
        user_id=data['user_id'],  # Assuming user_id is provided in the request
        project_name=data['project_name'],
        description=data['description'],
        url=data['url']
    )
    db.session.add(new_project)
    db.session.commit()
    return jsonify({'id': new_project.id}), 201




@app.route('/project-bullet-point/<int:bullet_point_id>', methods=['PUT'])
def update_bullet_point_project(bullet_point_id):
    bullet_point = ProjectBulletPoint.query.get(bullet_point_id)
    if not bullet_point or bullet_point.project.user_id != current_user.id:
        return jsonify({'message': 'Bullet point not found or access denied'}), 404

    data = request.json
    bullet_point.text = data['text']
    db.session.commit()
    return jsonify({'message': 'Bullet point updated'}), 200

@app.route('/project-bullet-point/<int:bullet_point_id>', methods=['DELETE'])
def delete_project_bullet_point(bullet_point_id):
    bullet_point = ProjectBulletPoint.query.get(bullet_point_id)
    if not bullet_point or bullet_point.project.user_id != current_user.id:
        return jsonify({'message': 'Bullet point not found or access denied'}), 404

    db.session.delete(bullet_point)
    db.session.commit()
    return jsonify({'message': 'Project Bullet Point deleted'}), 200


@app.route('/projects/<int:project_id>', methods=['PUT'])
def update_project_project(project_id):
    project = Projects.query.get(project_id)
    if not project or project.user_id != current_user.id:
        return jsonify({'message': 'Project not found or access denied'}), 404

    data = request.json
    project.project_name = data['project_name']
    project.description = data['description']
    project.url = data['url']

    db.session.commit()
    return jsonify({'message': 'Project updated'}), 200

@app.route('/projects/<int:project_id>', methods=['DELETE'])
def delete_project_project(project_id):
    project = Projects.query.get(project_id)
    if not project or project.user_id != current_user.id:
        return jsonify({'message': 'Project not found or access denied'}), 404

    # Delete associated bullet points
    for bp in project.bullet_points:
        db.session.delete(bp)

    db.session.delete(project)
    db.session.commit()
    return jsonify({'message': 'Project and associated bullet points deleted'}), 200

@app.route('/project-bullet-point', methods=['POST'])
def add_project_bullet_point():
    data = request.json
    user_id = current_user.id if current_user.is_authenticated else None

    project = Projects.query.filter_by(id=data['project_id'], user_id=user_id).first()
    if not project:
        return jsonify({'message': 'Project not found or access denied'}), 404

    new_bullet_point = ProjectBulletPoint(
        project_id=data['project_id'],
        text=data['text']
    )
    db.session.add(new_bullet_point)
    db.session.commit()
    return jsonify({'id': new_bullet_point.id}), 201

############ RESUME BUILDER - EDUCATION ###############

@app.route('/education', methods=['GET', 'POST'])
def education():
    if request.method == 'GET':
        user_id = current_user.id if current_user.is_authenticated else None
        if user_id is None:
            return jsonify([])  # Return an empty list if the user is not authenticated

        # Fetch education records specific to the logged-in user
        education_records = Education.query.filter_by(user_id=user_id).all()
        return jsonify([{
            'id': record.id,
            'degree': record.degree,
            'institution': record.institution,
            'graduation_year': record.graduation_year
        } for record in education_records])

    elif request.method == 'POST':
        if not current_user.is_authenticated:
            return jsonify({'message': 'Unauthorized'}), 401

        data = request.json

        new_education = Education(
            user_id=current_user.id,
            degree=data['degree'],
            institution=data['institution'],
            graduation_year=data['graduation_year']
        )
        db.session.add(new_education)
        db.session.commit()
        return jsonify({'id': new_education.id}), 201


@app.route('/education/<int:education_id>', methods=['PUT'])
def update_education(education_id):
    if not current_user.is_authenticated:
        return jsonify({'message': 'Unauthorized'}), 401

    education_record = Education.query.get(education_id)
    if not education_record or education_record.user_id != current_user.id:
        return jsonify({'message': 'Education record not found or access denied'}), 404

    data = request.json
    education_record.degree = data['degree']
    education_record.institution = data['institution']
    education_record.graduation_year = data['graduation_year']

    db.session.commit()
    return jsonify({'message': 'Education record updated'}), 200



@app.route('/education/<int:education_id>', methods=['DELETE'])
def delete_education(education_id):
    if not current_user.is_authenticated:
        return jsonify({'message': 'Unauthorized'}), 401

    education_record = Education.query.get(education_id)
    if not education_record or education_record.user_id != current_user.id:
        return jsonify({'message': 'Education record not found or access denied'}), 404

    db.session.delete(education_record)
    db.session.commit()
    return jsonify({'message': 'Education record deleted'}), 200



############ RESUME BUILDER - CERTIFICATIONS ###################################################

# Endpoint for fetching and adding certifications
@app.route('/certifications', methods=['GET', 'POST'])
def certifications():
    if request.method == 'GET':
        user_id = current_user.id if current_user.is_authenticated else None
        if user_id is None:
            return jsonify([])  # Return an empty list if the user is not authenticated

        # Fetch certification records specific to the logged-in user
        certification_records = Certifications.query.filter_by(user_id=user_id).all()
        return jsonify([{
            'id': record.id,
            'certification_name': record.certification_name,
            'issued_by': record.issued_by,
            'issue_date': record.issue_date.strftime('%Y-%m-%d') if record.issue_date else None,
            'expiration_date': record.expiration_date.strftime('%Y-%m-%d') if record.expiration_date else None
        } for record in certification_records])

    elif request.method == 'POST':
        if not current_user.is_authenticated:
            return jsonify({'message': 'Unauthorized'}), 401

        data = request.json

        new_certification = Certifications(
            user_id=current_user.id,
            certification_name=data['certification_name'],
            issued_by=data['issued_by'],
            issue_date=datetime.strptime(data['issue_date'], '%Y-%m-%d').date() if data['issue_date'] else None,
            expiration_date=datetime.strptime(data['expiration_date'], '%Y-%m-%d').date() if data['expiration_date'] else None,
        )
        db.session.add(new_certification)
        db.session.commit()
        return jsonify({'id': new_certification.id}), 201

# Endpoint for updating certification records
@app.route('/certifications/<int:certification_id>', methods=['PUT'])
def update_certification(certification_id):
    # Check if the user is authenticated
    if not current_user.is_authenticated:
        return jsonify({'message': 'Unauthorized'}), 401

    # Retrieve the certification record while ensuring it belongs to the current user
    certification_record = Certifications.query.filter_by(id=certification_id, user_id=current_user.id).first()
    if not certification_record:
        return jsonify({'message': 'Certification record not found or access denied'}), 404

    data = request.json
    certification_record.certification_name = data['certification_name']
    certification_record.issued_by = data['issued_by']
    certification_record.issue_date = datetime.strptime(data['issue_date'], '%Y-%m-%d').date() if data['issue_date'] else None
    certification_record.expiration_date = datetime.strptime(data['expiration_date'], '%Y-%m-%d').date() if data['expiration_date'] else None

    db.session.commit()
    return jsonify({'message': 'Certification record updated'}), 200

# Endpoint for deleting certification records
@app.route('/certifications/<int:certification_id>', methods=['DELETE'])
def delete_certification(certification_id):
    if not current_user.is_authenticated:
        return jsonify({'message': 'Unauthorized'}), 401

    certification_record = Certifications.query.get(certification_id)
    if not certification_record or certification_record.user_id != current_user.id:
        return jsonify({'message': 'Certification record not found or access denied'}), 404

    db.session.delete(certification_record)
    db.session.commit()
    return jsonify({'message': 'Certification record deleted'}), 200



############ RESUME BUILDER - PERSONAL INFORMATION ###################################################

# Endpoint for fetching and adding personal information
@app.route('/personal_information', methods=['GET', 'POST'])
def personal_information():
    if request.method == 'GET':
        user_id = current_user.id if current_user.is_authenticated else None
        if user_id is None:
            return jsonify([])  # Return an empty list if the user is not authenticated

        # Fetch personal information records specific to the logged-in user
        personal_info_records = PersonalInformation.query.filter_by(user_id=user_id).all()
        return jsonify([{
            'id': record.id,
            'full_name': record.full_name,
            'email': record.email,
            'contact_number': record.contact_number,
            'city_of_residence': record.city_of_residence,
            'state_of_residence': record.state_of_residence,
            'website': record.website,
            'github': record.github,
            'linkedin': record.linkedin
        } for record in personal_info_records])

    elif request.method == 'POST':
        if not current_user.is_authenticated:
            return jsonify({'message': 'Unauthorized'}), 401

        data = request.json

        new_personal_info = PersonalInformation(
            user_id=current_user.id,
            full_name=data['full_name'],
            email=data['email'],
            contact_number=data['contact_number'],
            city_of_residence=data['city_of_residence'],
            state_of_residence=data['state_of_residence'],
            website=data['website'],
            github=data['github'],
            linkedin=data['linkedin'],
        )
        db.session.add(new_personal_info)
        db.session.commit()
        return jsonify({'id': new_personal_info.id}), 201

# Endpoint for updating personal information records
@app.route('/personal_information/<int:personal_info_id>', methods=['PUT'])
def update_personal_information(personal_info_id):
    if not current_user.is_authenticated:
        return jsonify({'message': 'Unauthorized'}), 401

    personal_info_record = PersonalInformation.query.get(personal_info_id)
    if not personal_info_record or personal_info_record.user_id != current_user.id:
        return jsonify({'message': 'Personal information record not found or access denied'}), 404

    data = request.json
    personal_info_record.full_name = data['full_name']
    personal_info_record.email = data['email']
    personal_info_record.contact_number = data['contact_number']
    personal_info_record.city_of_residence = data['city_of_residence']
    personal_info_record.state_of_residence = data['state_of_residence']
    personal_info_record.website = data['website']
    personal_info_record.github = data['github']
    personal_info_record.linkedin = data['linkedin']

    db.session.commit()
    return jsonify({'message': 'Personal information record updated'}), 200

# Endpoint for deleting personal information records
@app.route('/personal_information/<int:personal_info_id>', methods=['DELETE'])
def delete_personal_information(personal_info_id):
    if not current_user.is_authenticated:
        return jsonify({'message': 'Unauthorized'}), 401

    personal_info_record = PersonalInformation.query.get(personal_info_id)
    if not personal_info_record or personal_info_record.user_id != current_user.id:
        return jsonify({'message': 'Personal information record not found or access denied'}), 404

    db.session.delete(personal_info_record)
    db.session.commit()
    return jsonify({'message': 'Personal information record deleted'}), 200



    ############ RESUME BUILDER - SKILLS ###################################################

# Endpoint for fetching and adding skills
@app.route('/skills', methods=['GET', 'POST'])
def skills():
    if request.method == 'GET':
        # Fetch all skills records for the current user
        user_id = current_user.id if current_user.is_authenticated else None
        if user_id is not None:
            skills_records = Skills.query.filter_by(user_id=user_id).all()
            return jsonify([{
                'id': record.id,
                'skill': record.skill
            } for record in skills_records])
        else:
            return jsonify([])  # Return an empty list if the user is not authenticated

    # Add new skill record
    data = request.json

    new_skill = Skills(
        user_id=current_user.id,  # Assuming the user is authenticated
        skill=data['skill']
    )
    db.session.add(new_skill)
    db.session.commit()
    return jsonify({'id': new_skill.id}), 201

# Endpoint for updating skill records
@app.route('/skills/<int:skill_id>', methods=['PUT'])
def update_skill(skill_id):
    skill_record = Skills.query.get(skill_id)
    if not skill_record:
        return jsonify({'message': 'Skill record not found'}), 404

    data = request.json
    skill_record.skill = data['skill']

    db.session.commit()
    return jsonify({'message': 'Skill record updated'}), 200

# Endpoint for deleting skill records
@app.route('/skills/<int:skill_id>', methods=['DELETE'])
def delete_skill(skill_id):
    skill_record = Skills.query.get(skill_id)
    if not skill_record:
        return jsonify({'message': 'Skill record not found'}), 404

    db.session.delete(skill_record)
    db.session.commit()
    return jsonify({'message': 'Skill record deleted'}), 200

@app.route('/generate-skills', methods=['POST'])
def generate_skills():
    # Check if user is authorized for either Tier 1 or Tier 2
    if not is_authorized_for_tier("Tier1") and not is_authorized_for_tier("Tier2"):
        return jsonify({"error": "Unauthorized access"}), 403
    data = request.json
    job_titles = data['job_titles']

    # Construct the prompt for OpenAI
    prompt = (f"List common skills for resumes with job titles: {', '.join(job_titles)}. "
              "The skills should be relevant and professionally useful for these roles.")

    try:
        completion = client.chat.completions.create(
            model="gpt-3.5-turbo-1106",
            messages=[
                {
                    "role": "user",
                    "content": prompt
                }
            ],
            temperature=0.7,  # Adjusted for creativity in generating diverse skills
            max_tokens=200,    # Adjusted to accommodate several skills
            top_p=1,
            frequency_penalty=0,
            presence_penalty=0
        )
        
        # Extract skills from response and process them
        skills_raw = completion.choices[0].message.content.strip().split('\n')
        skills_processed = [re.sub(r"^[0-9-.]+\s*", "", skill) for skill in skills_raw]

        return jsonify({'skills': skills_processed})

    except Exception as e:
        return jsonify({'error': str(e)}), 500


    ############ RESUME BUILDER - SUMMARY ###################################################

@app.route('/summary/<int:summary_id>', methods=['PUT'])
def update_summary(summary_id):
    if not current_user.is_authenticated:
        return jsonify({'message': 'Unauthorized'}), 401

    summary = Summary.query.get(summary_id)
    if not summary or summary.user_id != current_user.id:
        return jsonify({'message': 'Summary not found or access denied'}), 404

    data = request.json
    summary.summary_text = data['summary_text']

    try:
        db.session.commit()
        return jsonify({'message': 'Summary updated successfully!'})
    except Exception as e:
        db.session.rollback()
        return jsonify({'message': f'Error updating summary: {str(e)}'}), 500


@app.route('/summary', methods=['GET', 'POST'])
def handle_summary():
    if request.method == 'GET':
        if not current_user.is_authenticated:
            return jsonify({'message': 'Unauthorized'}), 401

        user_id = current_user.id
        summaries = Summary.query.filter_by(user_id=user_id).all()
        return jsonify([{'id': summary.id, 'summary_text': summary.summary_text} for summary in summaries])

    elif request.method == 'POST':
        if not current_user.is_authenticated:
            return jsonify({'message': 'Unauthorized'}), 401

        data = request.json
        summary_text = data['summary_text']

        # Check if a summary already exists for this user
        existing_summary = Summary.query.filter_by(user_id=current_user.id).first()
        if existing_summary:
            existing_summary.summary_text = summary_text
            db.session.commit()
            return jsonify({'message': 'Summary updated successfully!'}), 200

        new_summary = Summary(user_id=current_user.id, summary_text=summary_text)
        try:
            db.session.add(new_summary)
            db.session.commit()
            return jsonify({'message': 'Summary added successfully!'}), 201
        except Exception as e:
            db.session.rollback()
            return jsonify({'message': f'Error adding summary: {str(e)}'}), 500

@app.route('/summary/<int:summary_id>', methods=['DELETE'])
def delete_summary(summary_id):
    if not current_user.is_authenticated:
        return jsonify({'message': 'Unauthorized'}), 401

    summary = Summary.query.get(summary_id)
    if not summary or summary.user_id != current_user.id:
        return jsonify({'message': 'Summary not found or access denied'}), 404

    try:
        db.session.delete(summary)
        db.session.commit()
        return jsonify({'message': 'Summary deleted successfully!'}), 200
    except Exception as e:
        db.session.rollback()
        return jsonify({'message': f'Error deleting summary: {str(e)}'}), 500

############################## MY PROFILE ##########################################################################

# # # # # # # # # RESUME UPLOAD - DOCX ## # # # # # # # # # # # #

@app.route('/myprofile_current_resume')
def myprofile_current_resume():
    user_id = current_user.id if current_user.is_authenticated else None

    # Retrieve the most recent resume filename for the current user
    if user_id is not None:
        most_recent_resume = UserResumes.query.filter_by(user_id=user_id).order_by(UserResumes.uploaded_at.desc()).first()
        if most_recent_resume:
            most_recent_filename = most_recent_resume.filename
        else:
            most_recent_filename = None
    else:
        most_recent_filename = None

@app.route('/myprofile_resume_upload', methods=['POST'])
def myprofile_resume_upload():
    try:
        docx_file = request.files['docx_file']
        if docx_file and docx_file.filename.endswith('.docx'):
            # Get the user_id dynamically based on the currently authenticated user
            user_id = current_user.id if current_user.is_authenticated else None

            if user_id is None:
                return 'User not authenticated. Please log in.'

            # Read the DOCX file and extract text
            document = Document(BytesIO(docx_file.read()))
            resume_text = '\n'.join([para.text for para in document.paragraphs])

            filename = docx_file.filename

            # Delete the most recent resume
            most_recent_resume = UserResumes.query.filter_by(user_id=user_id).order_by(UserResumes.uploaded_at.desc()).first()
            if most_recent_resume:
                db.session.delete(most_recent_resume)

            # Add the new resume with extracted text
            new_resume = UserResumes(user_id, resume_text, filename)
            db.session.add(new_resume)
            db.session.commit()

            # Redirect to the /myprofile_current_resume route after successful upload
            return redirect(url_for('myprofile'))
        else:
            return 'Invalid file format. Please upload a DOCX file.'
    except Exception as e:
        return f'An error occurred: {str(e)}'


        # # # # ## # # DISPLAY RESUME NAME # # # ## # # #
@app.route('/myprofile')
def myprofile():
    user_id = current_user.id if current_user.is_authenticated else None

    # Retrieve the most recent resume filename for the current user
    if user_id is not None:
        most_recent_resume = UserResumes.query.filter_by(user_id=user_id).order_by(UserResumes.uploaded_at.desc()).first()
        if most_recent_resume:
            most_recent_filename = most_recent_resume.filename
        else:
            most_recent_filename = None
    else:
        most_recent_filename = None

    # Query the most recent UserDetails record
    most_recent_user_details = UserDetails.query.order_by(UserDetails.UserID.desc()).first()

    return render_template('myprofile.html', user_id=user_id, most_recent_filename=most_recent_filename, most_recent_user_details=most_recent_user_details)




        # # # # ## # # USER DETAILS CRUD ROUTES # # # ## # # #

@app.route('/delete_user_details', methods=['DELETE'])
def delete_user_details():
    # Retrieve user_id from the AJAX request's JSON body
    data = request.get_json()
    if not data or 'user_id' not in data:
        return jsonify({'success': False, 'message': 'User ID not provided'}), 400

    user_id = data['user_id']
    
    # Additional validation can be performed here if needed
    try:
        user_id = int(user_id)
    except ValueError:
        return jsonify({'success': False, 'message': 'Invalid user ID'}), 400

    user_details = UserDetails.query.get(user_id)
    if not user_details:
        return jsonify({'success': False, 'message': 'User not found'}), 404

    # Delete the UserDetails record from the database
    db.session.delete(user_details)
    db.session.commit()

    return jsonify({'success': True, 'message': 'UserDetails deleted successfully'})


@app.route('/update-user-details', methods=['POST'])
def update_user_details():
    if request.method == 'POST':
        # Get data from the form
        user_id = request.form.get('user_id')

        # Convert user_id to integer and handle potential errors
        try:
            user_id = int(user_id)
        except (TypeError, ValueError):
            return jsonify({'success': False, 'error': 'Invalid user ID format'}), 400

        # Fetch the specific user's details based on user_id
        user_details = UserDetails.query.get(user_id)
        if user_details:
            # Update user details
            user_details.FullName = request.form.get('full_name')
            user_details.RecentPosition = request.form.get('recent_position')
            user_details.DesiredJobTitle = request.form.get('desired_job_title')
            user_details.DesiredJobLocation = request.form.get('desired_job_location')
            user_details.DesiredWorkType = request.form.get('desired_work_type')
            user_details.DesiredCompensation = request.form.get('desired_compensation')
            user_details.JobAlertNotifications = request.form.get('job_alert_notifications')

            # Commit the changes to the database
            db.session.commit()
            return jsonify({'success': True, 'message': 'User details updated successfully'})
        else:
            return jsonify({'success': False, 'error': 'User not found'}), 404

    return jsonify({'success': False, 'error': 'Invalid request method'}), 400

@app.route('/create_user_details', methods=['POST'])
def create_user_details():
    if request.method == 'POST':
        user_id = request.form.get('user_id')  # Make sure this is being sent from the front-end
        full_name = request.form.get('full_name')
        recent_position = request.form.get('recent_position')
        desired_job_title = request.form.get('desired_job_title')
        desired_job_location = request.form.get('desired_job_location')
        desired_work_type = request.form.get('desired_work_type')
        desired_compensation = request.form.get('desired_compensation')
        job_alert_notifications = request.form.get('job_alert_notifications')

        if not user_id:
            return jsonify({'success': False, 'message': 'User ID is required'}), 400

        # Convert user_id to an integer
        try:
            user_id = int(user_id)
        except ValueError:
            return jsonify({'success': False, 'message': 'Invalid User ID'}), 400

        new_user_details = UserDetails(
            UserID=user_id,
            FullName=full_name,
            RecentPosition=recent_position,
            DesiredJobTitle=desired_job_title,
            DesiredJobLocation=desired_job_location,
            DesiredWorkType=desired_work_type,
            DesiredCompensation=desired_compensation,
            JobAlertNotifications=job_alert_notifications
        )

        db.session.add(new_user_details)
        db.session.commit()

        return jsonify({'success': True, 'message': 'User details created successfully'})

    return jsonify({'success': False, 'message': 'Invalid request method'}), 400

@app.route('/get-user-details')
def get_user_details():
    user_id = request.args.get('user_id')
    if user_id:
        try:
            user_id = int(user_id)
            user_details = UserDetails.query.get(user_id)
            if user_details:
                return jsonify({
                    'success': True,
                    'user': {
                        'full_name': user_details.FullName,
                        'recent_position': user_details.RecentPosition,
                        'desired_job_title': user_details.DesiredJobTitle,
                        'desired_job_location': user_details.DesiredJobLocation,
                        'desired_work_type': user_details.DesiredWorkType,
                        'desired_compensation': user_details.DesiredCompensation,
                        'job_alert_notifications': user_details.JobAlertNotifications
                    }
                })
            else:
                return jsonify({'success': False, 'error': 'User details not found'}), 404
        except ValueError:
            # Handle the case where user_id is not an integer
            return jsonify({'success': False, 'error': 'Invalid user ID format'}), 400
    else:
        return jsonify({'success': False, 'error': 'User ID not provided'}), 400



#@app.route('/success')
def success():
    # This route can be a success page where you inform the user that their details were updated successfully
    return "User details updated successfully"

@app.route('/error')
def error():
    # Handle errors and display an appropriate error message
    return 'An error occurred. Please try again.'

@app.route('/view-resume')
def view_resume():
    user_id = current_user.id if current_user.is_authenticated else None
    if user_id is None:
        # Handle unauthenticated case
        return "Please log in to view the resume", 401

    # Call the 'resume' function to get the complete resume data
    return resume(user_id)


@app.route("/download_resume_pdf", methods=["GET"])
def download_resume_pdf():
    user_id = current_user.id if current_user.is_authenticated else None
    if user_id is None:
        # Handle unauthenticated case
        return "User not authenticated", 401

    # Use the existing 'resume' function to get the complete resume data in HTML format
    # The 'resume' function should return an HTML string of the resume
    html_content = resume(user_id)

    # Convert the HTML to PDF using pdfkitd
    pdf_content = pdfkit.from_string(html_content, False)

    # Stream the PDF back to the user
    response = Response(stream_with_context(io.BytesIO(pdf_content)))
    response.headers["Content-Type"] = "application/pdf"
    response.headers["Content-Disposition"] = "inline"
    return response

@app.route('/rephrase-summary', methods=['POST'])
def process_summary():
       # Check if user is authorized for either Tier 1 or Tier 2
    if not is_authorized_for_tier("Tier1") and not is_authorized_for_tier("Tier2"):
        return jsonify({"error": "Unauthorized access"}), 403
    data = request.json
    summary_text = data['summary']

    # Revised prompt with specific instructions for output format
    prompt = (f"Please rephrase the following summary from a resume to make it sound more professional. "
              f"Provide the rephrased summary in a clear and concise manner, without any additional commentary or explanation.Very Important: Do not include any headers or titles before the revised summary. Your response should include just the rephrased summary and that is all. Do not include any text like 'Revised Summary:' before the summary.\n\n"
              f"Original Summary: {summary_text}")

    try:
        completion = client.chat.completions.create(
            model="gpt-3.5-turbo",
            messages=[
                {
                    "role": "user",
                    "content": prompt
                }
            ],
            temperature=0.5,  # Adjusted for more precise output
            max_tokens=150,    # Adjusted to limit the length of the response
            top_p=1,
            frequency_penalty=0,
            presence_penalty=0
        )
        
        # Accessing the response content
        rephrased_summary = completion.choices[0].message.content.strip()

        return jsonify({'rephrased_summary': rephrased_summary})

    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/extract-text', methods=['POST'])
def extract_text():
    if 'file' not in request.files:
        return jsonify({'error': 'No file part'}), 400

    file = request.files['file']
    if file.filename == '':
        return jsonify({'error': 'No selected file'}), 400

    if file and allowed_file(file.filename):
        doc = docx.Document(file)
        extracted_text = '\n'.join([para.text for para in doc.paragraphs])
        return jsonify({'extracted_text': extracted_text})

def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in {'docx'}


@app.route('/process-resume', methods=['POST'])
def process_resume():
    data = request.json
    resume_text = data.get('resume_text')

    # Read the HTML template content from a file
    template_path = Path(app.root_path) / 'templates' / 'ATS_score_template.html'
    with open(template_path, 'r', encoding='utf-8') as file:
        ats_score_template = file.read()

    # Construct the prompt for OpenAI
    prompt = (
        f"Please analyze the provided resume text for its ATS friendliness. Assess each section and give it a score (%). Provide feedback for each section, an overall summary of the analysis, a list of extracted keywords, and a bullet point list of action items for improvement. Use the provided html template to create an analysis report. Provide only the HTML code in your response.\n"
        f"Rules:\n"
        f"1. Your response should include only the HTML code, beginning with <html> and ending with </html>, and nothing more.\n"
        f"2. VERY IMPORTANT: Do not include any markdown formatting symbols like '```'. Exclude any additional text or explanations.\n"
        f"\nATS score template: {ats_score_template}"
        f"\nResume text: {resume_text}"
        f"\nSample response: <html>....</html>"
        f"\nVery important: DO NOT USE ANY MARKDOWN FORMATTING SYMBOLS LIKE '```'. EXAMPLE RESPONSE: <html>...</html>"
    )

    try:
        completion = client.chat.completions.create(
            model="gpt-4-1106-preview",  
            messages=[
                {
                    "role": "user",
                    "content": prompt
                }
            ],
            temperature=0.25,
            top_p=1,
            frequency_penalty=0,
            presence_penalty=0
        )
        
        # Get the response and strip Markdown code block symbols
        html_response = completion.choices[0].message.content.strip()
        if html_response.startswith("```html"):
            html_response = html_response[7:]  # Remove the starting ```html
        if html_response.endswith("```"):
            html_response = html_response[:-3]  # Remove the ending ```

        return jsonify({'html_response': html_response})

    except Exception as e:
        return jsonify({'error': str(e)}), 500


        # # # # ## # # PERSONALIZED JOB SEARCH # # # ## # # #
@app.route('/get-user-preferences')
def get_user_preferences():
    user_id = current_user.id if current_user.is_authenticated else None
    if user_id:
        user_details = UserDetails.query.filter_by(UserID=user_id).first()
        if user_details:
            return jsonify({
                'desired_job_title': user_details.DesiredJobTitle,
                'desired_job_location': user_details.DesiredJobLocation,
                'desired_work_type': user_details.DesiredWorkType
            })
    return jsonify({'error': 'User not authenticated or preferences not set'})


# # # # ## # # STRIPE CHECKOUT # # # ## # # #
def get_or_create_stripe_customer_id(email):
    with create_connection() as conn:
        with conn.cursor() as cursor:
            cursor.execute("SELECT stripe_customer_id FROM dbo.users WHERE email = ?", email)
            row = cursor.fetchone()
            if row and row.stripe_customer_id:
                return row.stripe_customer_id
            else:
                customer = stripe.Customer.create(email=email)
                stripe_customer_id = customer['id']
                cursor.execute("UPDATE dbo.users SET stripe_customer_id = ? WHERE email = ?", stripe_customer_id, email)
                conn.commit()
                return stripe_customer_id

@app.route('/create-checkout-session', methods=['POST'])
def create_checkout_session():
    try:
        user_email = request.form['email']
        stripe_customer_id = get_or_create_stripe_customer_id(user_email)

        price_id_mapping = {
            'Tier1': 'price_1OJOPqBmOXAq5RyDWYawmWvf',
            'Tier2': 'price_1OJOPyBmOXAq5RyDzIAxE0tc',
        }

        selected_plan = request.form['plan']
        price_id = price_id_mapping.get(selected_plan)

        if not price_id:
            return "Invalid plan selected", 400

        checkout_session = stripe.checkout.Session.create(
            customer=stripe_customer_id,
            payment_method_types=['card'],
            line_items=[{'price': price_id, 'quantity': 1}],
            mode='subscription',
            success_url='https://app.mycareermax.com/success?session_id={CHECKOUT_SESSION_ID}',
            cancel_url='https://app.mycareermax.com/cancel',
        )
        return redirect(checkout_session.url, code=303)
    except Exception as e:
        print(e)
        return jsonify(error="Server error", message=str(e)), 500

@app.route('/create-portal-session', methods=['POST'])
def customer_portal():
    # For demonstration purposes, we're using the Checkout session to retrieve the customer ID.
    # Typically this is stored alongside the authenticated user in your database.
    checkout_session_id = request.form.get('session_id')
    checkout_session = stripe.checkout.Session.retrieve(checkout_session_id)

    # This is the URL to which the customer will be redirected after they are
    # done managing their billing with the portal.
    return_url = 'http://app.mycareermax.com/dashboard'

    portalSession = stripe.billing_portal.Session.create(
        customer=checkout_session.customer,
        return_url=return_url,
    )
    return redirect(portalSession.url, code=303)


@app.route('/get-user-email')
def get_user_email():
    # Retrieve the username from the session
    username = session.get('username')
    if not username:
        return jsonify(success=False, error="User not logged in or username not found."), 401
    
    # Look up the email based on the username in the database
    try:
        with create_connection() as conn:
            with conn.cursor() as cursor:
                cursor.execute("SELECT email FROM dbo.users WHERE username = ?", username)
                row = cursor.fetchone()
                if row:
                    return jsonify(success=True, email=row.email)
                else:
                    return jsonify(success=False, error="User not found."), 404
    except Exception as e:
        return jsonify(success=False, error=str(e)), 500


# Function to create a Stripe customer
def create_stripe_customer(email):
    customer = stripe.Customer.create(email=email)
    return customer['id']

# Function to get or create Stripe customer ID
def get_or_create_stripe_customer_id(email):
    conn = create_connection()
    if conn is not None:
        with conn.cursor() as cursor:
            cursor.execute("SELECT stripe_customer_id FROM dbo.users WHERE email = ?", email)
            result = cursor.fetchone()
            if result and result[0]:
                return result[0]
            else:
                stripe_customer_id = create_stripe_customer(email)
                cursor.execute("UPDATE dbo.users SET stripe_customer_id = ? WHERE email = ?", stripe_customer_id, email)
                conn.commit()
                return stripe_customer_id
        conn.close()
    else:
        raise Exception("Database connection failed")



def is_authorized_for_tier(tier_required):
    if current_user.is_authenticated:
        subscription = current_user.subscription
        if subscription and subscription.is_active() and subscription.tier == tier_required:
            return True
    return False



if __name__ == "__main__":
    app.run(host="0.0.0.0", port=5000)

#if __name__ == "__main__":
 #   app.run(debug=True, host="0.0.0.0", port=5000)







